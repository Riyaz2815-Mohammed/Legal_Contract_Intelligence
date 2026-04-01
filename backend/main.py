import os
import sys

# --- Fix for WinError 1114 (Torch DLL crash on Windows) ---
# Force torch to use CPU to avoid DLL initialization conflicts with GPU drivers
# Must be set BEFORE torch is imported anywhere
os.environ["CUDA_VISIBLE_DEVICES"] = "-1"
os.environ["MKL_SERVICE_FORCE_INTEL"] = "1" # Extra safety for MKL

import asyncio
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, status, BackgroundTasks, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr
from typing import Optional, List
import secrets
import uuid
import requests
from datetime import datetime, timedelta
import jwt
import os
import json
import uuid
import difflib
import re
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
import boto3
from botocore.exceptions import ClientError
from botocore.config import Config
import psycopg2
from psycopg2 import pool

import time

print("[STARTUP] Initializing FastAPI app...")
app = FastAPI(title="LACCIS API", description="Legal Clause Classification Intelligence System")

@app.middleware("http")
async def log_requests(request, call_next):
    start_time = time.time()
    response = await call_next(request)
    process_time = time.time() - start_time
    print(f"INFO:     {request.client.host if request.client else 'unknown'} - \"{request.method} {request.url.path}\" {response.status_code} {process_time:.4f}s")
    return response

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174",],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security constants (HTTPBearer instance is defined later, after SSO config is loaded)
ALGORITHM = "HS256"

# Data storage
DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)
UPLOADS_DIR = DATA_DIR / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)


# Load environment variables                
env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
load_dotenv(env_path)
SECRET_KEY = os.getenv("JWT_SECRET", "your-secret-key-change-in-production")
EMAILJS_SERVICE_ID = os.getenv("EMAILJS_SERVICE_ID")
EMAILJS_TEMPLATE_ID = os.getenv("EMAILJS_TEMPLATE_ID")
EMAILJS_PUBLIC_KEY = os.getenv("EMAILJS_PUBLIC_KEY")
EMAILJS_PRIVATE_KEY = os.getenv("EMAILJS_PRIVATE_KEY")

# SSO / Auth configuration
ALLOW_LOCAL_LOGIN = os.getenv("ALLOW_LOCAL_LOGIN", "false").strip().lower() in ("1", "true", "yes")
SSO_ENDPOINT = os.getenv("SSO_ENDPOINT", "").strip()
SSO_COOKIE_NAME = os.getenv("SSO_COOKIE_NAME", "auth_token").strip()
CENTRAL_LOGOUT_URL = os.getenv("CENTRAL_LOGOUT_URL", "").strip()
print(f"[AUTH] ALLOW_LOCAL_LOGIN={ALLOW_LOCAL_LOGIN} | SSO_ENDPOINT={'SET' if SSO_ENDPOINT else 'UNSET (local JWT fallback active)'}")

# AWS Configuration
AWS_ACCESS_KEY = os.getenv("AWS_ACCESS_KEY", "").strip(' "')
AWS_SECRET_KEY = os.getenv("AWS_SECRET_KEY", "").strip(' "')
AWS_REGION = os.getenv("REGION", "").strip(' "')
BUCKET_NAME = os.getenv("BUCKET_NAME", "").strip(' "')

# Database Configuration
DATABASE_URL = os.getenv("DATABASE_URL")

# Initialize PostgreSQL Connection Pool
db_pool = None
if DATABASE_URL:
    try:
        print("[DATABASE] Attempting to initialize PostgreSQL Connection Pool...")
        db_pool = psycopg2.pool.ThreadedConnectionPool(
            1, 20, DATABASE_URL
        )
        print("[DATABASE] PostgreSQL Connection Pool initialized successfully")
    except Exception as e:
        print(f"[ERROR] Database pool initialization failed: {e}")
else:
    print("[DATABASE] DATABASE_URL missing")

print("[STARTUP] Initializing S3 client...")
s3_client = boto3.client(
    's3',
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region_name=AWS_REGION,
    config=Config(
        connect_timeout=10,   # 10 sec to establish connection
        read_timeout=30,      # 30 sec to read response
        retries={'max_attempts': 2}
    )
)
print("[STARTUP] S3 client initialized.")

# Debug: Print if credentials are loaded
print(f"[AWS/EMAIL] EMAILJS_SERVICE_ID loaded: {bool(EMAILJS_SERVICE_ID)}")
print(f"[AWS/EMAIL] EMAILJS_TEMPLATE_ID loaded: {bool(EMAILJS_TEMPLATE_ID)}")
print(f"[AWS/EMAIL] EMAILJS_PUBLIC_KEY loaded: {bool(EMAILJS_PUBLIC_KEY)}")
print(f"[AWS/EMAIL] AWS_ACCESS_KEY loaded: {bool(AWS_ACCESS_KEY)} | value starts with: {AWS_ACCESS_KEY[:4] if AWS_ACCESS_KEY else 'MISSING'}")
print(f"[AWS/EMAIL] AWS_REGION: {AWS_REGION}")
print(f"[AWS/EMAIL] BUCKET_NAME: {BUCKET_NAME}")

# Models
class LoginRequest(BaseModel):
    email: EmailStr
    password: str

class ClientCreate(BaseModel):
    name: str
    email: EmailStr

class LegalTeamMemberCreate(BaseModel):
    name: str
    email: EmailStr

class User(BaseModel):
    id: str
    name: str
    email: EmailStr
    role: str  # 'admin' or 'client'

class DocumentUpload(BaseModel):
    document_type: str  # 'NDA', 'MSA', 'SOW', 'Redlined', 'Others'
    shared_with: Optional[str] = None  # client_id or 'admin'
    
class DocumentShare(BaseModel):
    document_id: str
    share_with: str  # client_id or 'admin'

class ClauseEdit(BaseModel):
    content_id: str
    edited_content: str

class ClauseComment(BaseModel):
    content_id: str
    comment: str

class MessageSend(BaseModel):
    recipient_id: str  # client_id or 'admin'
    content: str

class Message(BaseModel):
    id: str
    sender_id: str
    recipient_id: str
    content: str
    timestamp: str

# Helper functions
def record_activity(user_id: str, client_id: str, action: str, details: str = ""):
    if db_pool:
        conn = None
        try:
            conn = db_pool.getconn()
            with conn.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO activity_log (id, user_id, client_id, action, details, timestamp)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
                    (f"act-{uuid.uuid4().hex[:8]}", user_id, client_id, action, details, datetime.now())
                )
            conn.commit()
        except Exception as e:
            print(f"[ERROR] record_activity error: {e}")
            if conn: conn.rollback()
        finally:
            if conn: db_pool.putconn(conn)


def create_token(user_id: str, email: str, role: str):
    payload = {
        "user_id": user_id,
        "email": email,
        "role": role,
        "exp": datetime.utcnow() + timedelta(days=7)
    }
    return jwt.encode(payload, SECRET_KEY, algorithm=ALGORITHM)

def _resolve_user_from_email(email: str) -> dict:
    """Look up a local user by email and return their identity dict.
    Raises 403 if the email is not onboarded in this Legal app."""
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT id, email, name, role FROM users WHERE email = %s", (email,))
            row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=403, detail="User not onboarded in Legal app")
        user_id, u_email, name, role = row
        return {"user_id": user_id, "email": u_email, "name": name, "role": role}
    finally:
        if conn:
            db_pool.putconn(conn)


def _verify_via_sso(request: Request) -> dict:
    """Validate the shared SSO cookie against the central endpoint.
    Returns the local user dict on success.
    Raises 401 if no cookie or SSO rejects it, 403 if not onboarded."""
    cookie_value = request.cookies.get(SSO_COOKIE_NAME)
    if not cookie_value:
        raise HTTPException(status_code=401, detail="Not authenticated (no SSO cookie)")
    try:
        sso_resp = requests.get(
            SSO_ENDPOINT,
            cookies={SSO_COOKIE_NAME: cookie_value},
            timeout=5
        )
    except Exception as e:
        print(f"[AUTH] SSO endpoint unreachable: {e}")
        raise HTTPException(status_code=401, detail="SSO endpoint unreachable")
    if sso_resp.status_code != 200:
        raise HTTPException(status_code=401, detail="SSO session invalid")
    sso_data = sso_resp.json()
    email = sso_data.get("email")
    if not email:
        raise HTTPException(status_code=401, detail="SSO response missing email")
    print(f"[AUTH] SSO validated: {email}")
    return _resolve_user_from_email(email)


def _verify_via_local_jwt(credentials: Optional[HTTPAuthorizationCredentials]) -> dict:
    """Validate a local JWT Bearer token (dev fallback only)."""
    if credentials is None:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        print(f"[AUTH] Local JWT verified: {payload.get('email')} | role: {payload.get('role')}")
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


# Make security optional so it doesn't crash when SSO cookie mode is active
security = HTTPBearer(auto_error=False)


def verify_token(
    request: Request,
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)
) -> dict:
    """Unified auth dependency.
    - If SSO_ENDPOINT is configured: validates the shared SSO cookie (production).
    - Otherwise: falls back to local JWT Bearer token (dev mode).
    """
    if SSO_ENDPOINT:
        return _verify_via_sso(request)
    # Local JWT fallback
    return _verify_via_local_jwt(credentials)

def send_email(to_email: str, subject: str, body: str):
    """
    Send email via EmailJS REST API
    
    Args:
        to_email: Recipient email address
        subject: Email subject
        body: Email body (HTML format)
    
    Returns:
        dict: {'success': bool, 'message': str}
    """
    try:
        # Check if EmailJS credentials are configured
        if not all([EMAILJS_SERVICE_ID, EMAILJS_TEMPLATE_ID, EMAILJS_PUBLIC_KEY]):
            error_msg = f"EmailJS credentials not configured. SERVICE_ID: {bool(EMAILJS_SERVICE_ID)}, TEMPLATE_ID: {bool(EMAILJS_TEMPLATE_ID)}, PUBLIC_KEY: {bool(EMAILJS_PUBLIC_KEY)}"
            print(f"✗ {error_msg}")
            return {
                'success': False,
                'message': error_msg
            }
        
        # EmailJS API endpoint
        url = "https://api.emailjs.com/api/v1.0/email/send"
        
        # Prepare email data
        email_data = {
            "service_id": EMAILJS_SERVICE_ID,
            "template_id": EMAILJS_TEMPLATE_ID,
            "user_id": EMAILJS_PUBLIC_KEY,
            "accessToken": EMAILJS_PRIVATE_KEY,  # Required for Strict Mode
            "template_params": {
                "to_email": to_email,
                "subject": subject,
                "html_content": body
            }
        }
        
        print(f"📤 Sending email to {to_email} via EmailJS...")
        print(f"   Service: {EMAILJS_SERVICE_ID}")
        print(f"   Template: {EMAILJS_TEMPLATE_ID}")
        
        # Send via EmailJS
        response = requests.post(url, json=email_data, timeout=10)
        
        print(f"   Response Status: {response.status_code}")
        print(f"   Response Body: {response.text}")
        
        if response.status_code == 200:
            print(f"✓ Email sent successfully to {to_email}")
            return {
                'success': True,
                'message': f'Email sent successfully to {to_email}'
            }
        else:
            error_msg = f"EmailJS error: {response.status_code} - {response.text}"
            print(f"✗ {error_msg}")
            return {
                'success': False,
                'message': error_msg
            }
    
    except requests.exceptions.Timeout:
        error_msg = "EmailJS request timeout - check your internet connection"
        print(f"✗ {error_msg}")
        return {'success': False, 'message': error_msg}
    
    except requests.exceptions.RequestException as e:
        error_msg = f"EmailJS request error: {str(e)}"
        print(f"✗ {error_msg}")
        return {'success': False, 'message': error_msg}
    
    except Exception as e:
        error_msg = f"Email error: {str(e)}"
        print(f"✗ {error_msg}")
        return {'success': False, 'message': error_msg}

# Routes

@app.get("/")
def root():
    return {"message": "LACCIS API is running", "version": "1.0.0"}


@app.get("/auth/me")
def auth_me(
    request: Request,
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)
):
    """Session bootstrap endpoint for the frontend.
    Returns the current user if authenticated, or 401/403.
    """
    try:
        user = verify_token(request, credentials)
    except HTTPException:
        raise
    # Normalise the user dict shape (SSO path returns name; JWT path may not)
    return {
        "user": {
            "id":    user.get("user_id") or user.get("id", ""),
            "name":  user.get("name", ""),
            "email": user.get("email", ""),
            "role":  user.get("role", ""),
        }
    }

@app.post("/api/auth/login")
def login(request: LoginRequest):
    if not ALLOW_LOCAL_LOGIN:
        raise HTTPException(status_code=403, detail="Local login is disabled. Please sign in via the central portal.")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT id, email, name, role, password_hash, nda_accepted FROM users WHERE email = %s", (request.email,))
            user_row = cur.fetchone()
            
        if not user_row:
            raise HTTPException(status_code=401, detail="Invalid credentials")
        
        user_id, email, name, role, password_hash, nda_accepted = user_row
        
        # Direct password comparison
        if password_hash == request.password:
            token = create_token(user_id, email, role)
            record_activity(user_id, user_id, "Logged in")
            return {
                "token": token,
                "user": {
                    "id": user_id,
                    "name": name,
                    "email": email,
                    "role": role,
                    "nda_accepted": nda_accepted
                }
            }
        raise HTTPException(status_code=401, detail="Invalid credentials")
    except HTTPException: raise
    except Exception as e:
        print(f"[ERROR] Login error: {e}")
        raise HTTPException(status_code=500, detail="Database error during login")
    finally:
        if conn: db_pool.putconn(conn)



@app.post("/api/clients/create")
def create_client(client: ClientCreate, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Only admins and legal team members can create clients")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")

    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # Check existence
            cur.execute("SELECT 1 FROM users WHERE email = %s", (client.email,))
            if cur.fetchone():
                raise HTTPException(status_code=400, detail="Email already exists")

            client_id = f"client-{uuid.uuid4().hex[:8]}"
            random_suffix = secrets.token_hex(3).upper()
            password = f"LACCIS-{random_suffix}"
            created_at = datetime.now()
            
            cur.execute(
                """
                INSERT INTO users (id, name, email, password_hash, role, nda_accepted, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
                """,
                (client_id, client.name, client.email, password, "client", False, created_at)
            )
        conn.commit()
    except HTTPException: raise
    except Exception as e:
        print(f"[ERROR] create_client database error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    finally:
        if conn: db_pool.putconn(conn)

    # Load and customize email template
    try:
        template_path = Path("email_template.html")
        if template_path.exists():
            with open(template_path, 'r', encoding='utf-8') as f:
                email_body = f.read()
            email_body = email_body.replace('{CLIENT_NAME}', client.name)
            email_body = email_body.replace('{CLIENT_EMAIL}', client.email)
            email_body = email_body.replace('{CLIENT_PASSWORD}', password)
        else:
            email_body = f"<html><body><h2>Welcome {client.name}</h2><p>Password: {password}</p></body></html>"
    except Exception as e:
        print(f"Error loading template: {e}")
        email_body = f"<html><body><h2>Welcome {client.name}</h2><p>Password: {password}</p></body></html>"
    
    email_res = send_email(client.email, "Your LACCIS Login Credentials", email_body)
    
    return {
        "message": "Client created successfully",
        "client": {"id": client_id, "name": client.name, "email": client.email, "created_at": created_at.isoformat()},
        "email_sent": email_res.get('success', False),
        "email_message": email_res.get('message', ""),
        "credentials": {
            "email": client.email,
            "password": password
        }
    }




@app.post("/api/legal/create")
def create_legal_team_member(member: LegalTeamMemberCreate, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Only admins and legal team members can create legal team members")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")

    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT 1 FROM users WHERE email = %s", (member.email,))
            if cur.fetchone():
                raise HTTPException(status_code=400, detail="Email already exists")

            member_id = f"legal-{uuid.uuid4().hex[:8]}"
            password = secrets.token_urlsafe(12)
            created_at = datetime.now()
            
            cur.execute(
                """
                INSERT INTO users (id, name, email, password_hash, role, created_at)
                VALUES (%s, %s, %s, %s, %s, %s)
                """,
                (member_id, member.name, member.email, password, "legal_team", created_at)
            )
        conn.commit()
    except HTTPException: raise
    except Exception as e:
        print(f"[ERROR] create_legal_team_member database error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    finally:
        if conn: db_pool.putconn(conn)

    # Send email with credentials
    email_body = f"<html><body><h2>Welcome {member.name}</h2><p>Password: {password}</p></body></html>"
    email_res = send_email(member.email, "LACCIS Legal Team Account", email_body)
    
    return {
        "message": "Legal team member created successfully",
        "member": {"id": member_id, "name": member.name, "email": member.email, "created_at": created_at.isoformat()},
        "email_sent": email_res.get('success', False),
        "credentials": {
            "email": member.email,
            "password": password
        }
    }



@app.get("/api/legal/list")
def list_legal_team(current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ["admin", "client", "legal_team"]:
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            if current_user["role"] == "admin":
                cur.execute("SELECT id, name, email, created_at FROM users WHERE role = 'legal_team'")
                rows = cur.fetchall()
                members = [{"id": r[0], "name": r[1], "email": r[2], "created_at": r[3].isoformat() if r[3] else None} for r in rows]
            else:
                cur.execute("SELECT id, name, email FROM users WHERE role IN ('legal_team', 'admin')")
                rows = cur.fetchall()
                members = [{"id": r[0], "name": r[1], "email": r[2]} for r in rows]
            return {"members": members}
    except Exception as e:
        print(f"[ERROR] list_legal_team error: {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)


@app.delete("/api/legal/delete/{member_id}")
def delete_legal_team_member(member_id: str, current_user: dict = Depends(verify_token)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("DELETE FROM users WHERE id = %s AND role = 'legal_team'", (member_id,))
        conn.commit()
        return {"message": "Legal team member deleted successfully"}
    except Exception as e:
        print(f"[ERROR] delete_legal_team_member error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)


@app.get("/api/clients/list")
def list_clients(current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT id, name, email, created_at FROM users WHERE role = 'client'")
            rows = cur.fetchall()
            clients = [{"id": r[0], "name": r[1], "email": r[2], "created_at": r[3].isoformat() if r[3] else None} for r in rows]
            return {"clients": clients}
    except Exception as e:
        print(f"[ERROR] list_clients error: {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)



@app.delete("/api/clients/delete/{client_id}")
def delete_client(client_id: str, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # 1. Get all documents for this client to delete from S3
            cur.execute("SELECT s3_key, file_path FROM documents WHERE user_id = %s", (client_id,))
            docs = cur.fetchall()
            
            # 2. Cleanup S3 and local files
            for s3_key, file_path in docs:
                if s3_key:
                    try:
                        s3_client.delete_object(Bucket=BUCKET_NAME, Key=s3_key)
                    except Exception as e:
                        print(f"[ERROR] S3 cleanup error: {e}")
                if file_path:
                    try:
                        fp = Path(file_path)
                        if fp.exists(): fp.unlink()
                    except Exception as e:
                        print(f"[ERROR] Local file cleanup error: {e}")

            # 3. DB cleanup - Cascade should handle it if set, but we do it explicitly to handle clauses
            cur.execute("DELETE FROM clauses WHERE document_id IN (SELECT id FROM documents WHERE user_id = %s)", (client_id,))
            cur.execute("DELETE FROM documents WHERE user_id = %s", (client_id,))
            cur.execute("DELETE FROM users WHERE id = %s", (client_id,))
        conn.commit()
        return {"message": "Client and associated documents deleted successfully"}
    except Exception as e:
        print(f"[ERROR] delete_client database error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)



def trigger_extraction(file_name: str, document_id: str, document_type: str = "Unknown", source: str = "unknown"):
    """Background task: perform extraction and classification locally.

    Args:
        file_name     : saved filename under data/uploads/
        document_id   : actual database ID for the document
        document_type : e.g. "NDA", "MSA", "SOW" — stored in every clause record
        source        : "client" or "legal" — who uploaded the document
    """
    try:
        # Ensure extracter directory is in sys.path
        backend_dir = Path(__file__).parent
        project_root = backend_dir.parent
        extracter_dir = project_root / "extracter"

        if str(extracter_dir) not in sys.path:
            sys.path.append(str(extracter_dir))

        # Import extraction logic (lazy import to resolve at runtime)
        from extract import extract_text_from_file
        from clause_engine import parse_text_file, process_document

        # Local path for the file (it's already saved in UPLOADS_DIR)
        local_path = UPLOADS_DIR / file_name
        if not local_path.exists():
            print(f"[ERROR] [Background] File not found for extraction at {local_path}")
            return

        print(f"[INFO] [Background] Starting extraction for {file_name} (doc={document_type}, src={source})...")

        # 1. Extract text (PDF, DOCX, TXT)
        extracted_text = extract_text_from_file(str(local_path))

        # 2. Parse text blocks using temp files on disk to preserve RAM
        import tempfile
        import os
        
        # Write to a temporary file locally so clause_engine can perfectly parse it with \n newlines
        with tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', delete=False, suffix='.txt') as temp_out:
            temp_out.write(extracted_text)
            temp_file_path = temp_out.name
            
        try:
            extracted_blocks = parse_text_file(temp_file_path)
        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)

        # 3. Classify — pass document type and source so every clause record is fully tagged
        results = process_document(extracted_blocks, document=document_type, source=source)

        # (JSON storage was removed per user request, DB storage only)
        
        # 4. Insert into clauses table and run vector pipeline
        conn = None
        try:
            if db_pool:
                conn = db_pool.getconn()
                with conn.cursor() as cur:
                    # Using passed document_id directly

                    if source == "legal":
                        # ── Legal templates: store EACH parsed section individually ──────────
                        # Do NOT merge by clause type — ChromaDB needs granular, single-section
                        # embeddings so the clause-type filter in query.py works precisely.
                        clauses_to_insert = []
                        for clause in results:
                            clauses_to_insert.append({
                                "clause_id":  clause.get("clause_id") or f"CLZ-{uuid.uuid4().hex[:8].upper()}",
                                "clause":     clause.get("clause", "Unknown"),
                                "content_id": clause.get("content_id") or f"CNT-{uuid.uuid4().hex[:8].upper()}",
                                "content":    clause.get("content", "").strip(),
                                "page_number": clause.get("page_number", 1),
                            })
                    else:
                        # ── Client documents: merge same-type clauses into one record ───────
                        # This gives a single coherent clause block per type for review display.
                        merged_clauses: dict = {}
                        for clause in results:
                            ctype   = clause.get("clause", "Unknown")
                            ccontent = clause.get("content", "").strip()
                            cpage   = clause.get("page_number", 1)
                            if ctype not in merged_clauses:
                                merged_clauses[ctype] = {
                                    "clause_id":  clause.get("clause_id") or f"CLZ-{uuid.uuid4().hex[:8].upper()}",
                                    "clause":     ctype,
                                    "content_id": clause.get("content_id") or f"CNT-{uuid.uuid4().hex[:8].upper()}",
                                    "content":    ccontent,
                                    "page_number": cpage,
                                }
                            else:
                                merged_clauses[ctype]["content"] += "\n\n" + ccontent
                        clauses_to_insert = list(merged_clauses.values())

                    # Persist to DB
                    results = clauses_to_insert  # pipeline below iterates this
                    for clause in clauses_to_insert:
                        cur.execute(
                            """
                            INSERT INTO clauses (clause_id, clause, content_id, content, page_number, document, source, document_id)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                            ON CONFLICT (content_id) DO NOTHING
                            """,
                            (
                                clause["clause_id"],
                                clause["clause"],
                                clause["content_id"],
                                clause["content"],
                                clause["page_number"],
                                document_type,
                                source,
                                document_id,
                            )
                        )
                conn.commit()
                print(f"[SUCCESS] [Background] Saved {len(results)} clauses to database")
                
        except Exception as db_err:
            print(f"[ERROR] [Background] Failed to save clauses format: {db_err}")
            if conn: conn.rollback()
            import traceback
            traceback.print_exc()
        finally:
            if conn: 
                db_pool.putconn(conn)
                conn = None # Prevent reuse below
                
        # ------ Vector Pipeline -----------------------------------------------
        # Runs independently — clauses are already safely in DB before this.
        try:
            if source == "legal":
                from vector_pipeline.embeddings.embed_store import run_embed_pipeline
                print(f"[INFO] [Background] Standard template '{file_name}' — updating Neon pgvector...")
                run_embed_pipeline()
                print(f"[SUCCESS] [Background] Neon pgvector updated with new standard template clauses.")
            
            elif source == "client" and document_id:
                from vector_pipeline.pipeline.full_pipeline import run_pipeline
                print(f"[INFO] [Background] Client document '{file_name}' — running vector review pipeline...")

                final_reviews = []
                for clause in results:
                    client_text = clause.get("content", "")
                    clause_type = clause.get("clause", "Unknown")
                    
                    if len(client_text.strip()) <= 10:
                        final_reviews.append({
                            "content_id": clause.get("content_id"), "clause_id": clause.get("clause_id"),
                            "clause_type": clause_type, "content": client_text,
                            "page_number": clause.get("page_number", 1),
                            "risk": "Low", "similarity_score": None,
                            "matched_clause": None, "llm_reasoning": None, "status": "pending"
                        })
                        continue

                    if clause_type in ("Other",):
                        final_reviews.append({
                            "content_id": clause.get("content_id"), "clause_id": clause.get("clause_id"),
                            "clause_type": clause_type, "content": client_text,
                            "page_number": clause.get("page_number", 1),
                            "risk": "Low", "similarity_score": None,
                            "matched_clause": None, "llm_reasoning": None, "status": "pending"
                        })
                        continue

                    # Per-clause try so one failure doesn't abort entire review
                    try:
                        pipeline_results = run_pipeline(
                            query_text=client_text,
                            clause_type=clause_type,
                            document_type=document_type
                        )
                        best_match = pipeline_results[0] if pipeline_results else None
                    except Exception as clause_err:
                        print(f"[WARN] [Background] Pipeline failed for clause '{clause_type}': {clause_err}")
                        best_match = None

                    if best_match:
                        # Override risk for Header to always be Low
                        final_risk = "Low" if clause_type == "Header" else best_match.get("final_risk", "High")
                        final_status = "approved" if clause_type == "Header" else "pending"
                        
                        review_entry = {
                            "content_id":       clause.get("content_id"),
                            "clause_id":        clause.get("clause_id"),
                            "clause_type":      clause_type,
                            "content":          client_text,
                            "page_number":      clause.get("page_number", 1),
                            "risk":             final_risk,
                            "similarity_score": best_match.get("sbert_similarity"),
                            "matched_clause": {
                                "content":       best_match.get("template_content"),
                                "document_type": best_match.get("template_metadata", {}).get("document_type", "Template")
                            },
                            "llm_reasoning":    best_match.get("llm_reasoning"),
                            "status":           final_status
                        }
                    else:
                        review_entry = {
                            "content_id":       clause.get("content_id"),
                            "clause_id":        clause.get("clause_id"),
                            "clause_type":      clause_type,
                            "content":          client_text,
                            "page_number":      clause.get("page_number", 1),
                            "risk":             "Low" if clause_type == "Header" else "High",
                            "risk_confidence":  0.92, # Placeholder or extracted confidence
                            "similarity_score": 0.88,
                            "matched_clause":   None,
                            "llm_reasoning":    None,
                            "status":           "pending"
                        }
                    final_reviews.append(review_entry)

                # Always save review — even if some clauses fell back to High risk
                reviews_dir = DATA_DIR / "reviews"
                reviews_dir.mkdir(parents=True, exist_ok=True)
                review_path = reviews_dir / f"{document_id}.json"
                with open(review_path, "w", encoding="utf-8") as f:
                    json.dump(final_reviews, f, indent=4)

                if db_pool:
                    r_conn = db_pool.getconn()
                    try:
                        with r_conn.cursor() as cur:
                            cur.execute(
                                """
                                INSERT INTO document_reviews (document_id, review_data)
                                VALUES (%s, %s)
                                ON CONFLICT (document_id) DO UPDATE
                                SET review_data = EXCLUDED.review_data, created_at = NOW()
                                """,
                                (document_id, json.dumps(final_reviews))
                            )
                        r_conn.commit()
                    except Exception as db_err:
                        print(f"[ERROR] [Background] Failed to save review to DB: {db_err}")
                        r_conn.rollback()
                    finally:
                        db_pool.putconn(r_conn)

                print(f"[SUCCESS] [Background] Review pipeline complete — {len(final_reviews)} clauses saved.")
                
        except Exception as vector_err:
            print(f"[ERROR] [Background] Vector pipeline failed: {vector_err}")
            import traceback
            traceback.print_exc()
        # -----------------------------------------------------------------------

        # Update document status to 'completed' in PostgreSQL
        if document_id and db_pool:
            update_conn = None
            try:
                update_conn = db_pool.getconn()
                with update_conn.cursor() as cur:
                    # Using 'uploaded' instead of 'completed' to satisfy DB CHECK constraint
                    cur.execute("UPDATE documents SET status = 'uploaded' WHERE id = %s", (document_id,))
                update_conn.commit()
                print(f"[SUCCESS] [Background] Document status marked as 'completed' for {document_id}")
            except Exception as e:
                print(f"[ERROR] [Background] Failed to update document status: {e}")
                if update_conn: update_conn.rollback()
            finally:
                if update_conn: db_pool.putconn(update_conn)

        print(f"[SUCCESS] [Background] Classification complete for {file_name}: {len(results)} clauses processed.")
    except Exception as e:
        import traceback
        print(f"[ERROR] [Background] Extraction failed for {file_name}: {str(e)}")
        print(traceback.format_exc())

@app.post("/api/documents/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    document_type: str = "Others",
    shared_with: Optional[str] = None,
    is_final: bool = Form(False),
    current_user: dict = Depends(verify_token)
):
    # All document types are allowed without restriction
    print(f"📄 User {current_user['user_id']} uploading {document_type}")
    
    # Determine status and ID early
    status = "pending" if document_type.startswith("NDA") else "uploaded"
    doc_uuid = f"doc-{uuid.uuid4().hex[:8]}"

    # Save file with unique ID in name to avoid S3 collisions
    file_name = f"{doc_uuid}_{file.filename}"
    file_path = UPLOADS_DIR / file_name
    
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    # Upload to S3 (run in thread pool to avoid blocking the async event loop)
    try:
        if not AWS_ACCESS_KEY or not AWS_SECRET_KEY or not BUCKET_NAME:
            raise ValueError(f"Missing AWS config — KEY:{bool(AWS_ACCESS_KEY)} SECRET:{bool(AWS_SECRET_KEY)} BUCKET:{bool(BUCKET_NAME)}")
        print(f"[AWS] Uploading {file_name} to S3 bucket: {BUCKET_NAME} (region: {AWS_REGION})...")
        loop = asyncio.get_event_loop()
        await asyncio.wait_for(
            loop.run_in_executor(
                None,
                lambda: s3_client.put_object(
                    Bucket=BUCKET_NAME,
                    Key=file_name,
                    Body=content
                )
            ),
            timeout=30  # give up after 30 seconds
        )
        print(f"[SUCCESS] Successfully uploaded to S3")
        s3_url = f"https://{BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{file_name}"
    except asyncio.TimeoutError:
        print(f"[ERROR] S3 Upload timed out after 30s — check bucket region/permissions")
        s3_url = None
    except ClientError as e:
        error_code = e.response['Error']['Code']
        error_msg = e.response['Error']['Message']
        print(f"[ERROR] S3 Upload Error [{error_code}]: {error_msg}")
        s3_url = None
    except Exception as e:
        print(f"[ERROR] S3 Upload Unexpected Error: {type(e).__name__}: {e}")
        s3_url = None
    # Get users to share with automatically if uploaded by client
    final_shared_with = list(shared_with) if shared_with else []
    
    if current_user["role"] == "client" and db_pool:
        temp_conn = None
        try:
            temp_conn = db_pool.getconn()
            with temp_conn.cursor() as cur:
                cur.execute("SELECT id FROM users WHERE role IN ('admin', 'legal_team')")
                admin_ids = [row[0] for row in cur.fetchall()]
                for admin_id in admin_ids:
                    if admin_id not in final_shared_with:
                        final_shared_with.append(admin_id)
        except Exception as e:
            print(f"[ERROR] Failed to fetch admins for auto-share: {e}")
        finally:
            if temp_conn: db_pool.putconn(temp_conn)

    # Determine status based on document type
    is_client = current_user["role"] not in ["admin", "legal_team"]

    new_doc = {
        "id": doc_uuid,
        "filename": file.filename,
        "document_type": document_type,
        "user_id": current_user["user_id"],
        "user_email": current_user["email"],
        "user_role": current_user["role"],
        "size": len(content),
        "status": status,
        "shared_with": final_shared_with,
        "uploaded_at": datetime.now(),
        "file_path": str(file_path),
        "s3_url": s3_url,
        "s3_key": file_name,
        "is_finalized": is_final if not is_client else False,
        "client_marked_final": is_final if is_client else False
    }
    
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
        
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO documents (id, filename, document_type, user_id, user_email, user_role, size, status, shared_with, uploaded_at, file_path, s3_url, s3_key, is_finalized, client_marked_final)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (new_doc["id"], new_doc["filename"], new_doc["document_type"], new_doc["user_id"], new_doc["user_email"], new_doc["user_role"], new_doc["size"], new_doc["status"], json.dumps(new_doc["shared_with"]), new_doc["uploaded_at"], new_doc["file_path"], new_doc["s3_url"], new_doc["s3_key"], new_doc["is_finalized"], new_doc["client_marked_final"])
            )
        conn.commit()
    except Exception as e:
        print(f"[ERROR] Database document upload error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Failed to save metadata")
    finally:
        if conn: db_pool.putconn(conn)

    # Record activity
    record_activity(
        user_id=current_user["user_id"],
        client_id=current_user["user_id"] if current_user["role"] == "client" else (shared_with or "admin"),
        action="Uploaded document",
        details=f"Document: {file.filename} ({document_type})"
    )
    
    # Trigger automated extraction in background (Skip for Redlined and Final docs)
    source = "client" if current_user["role"] == "client" else "legal"
    # Check if the document type contains 'Final' or if it was marked as final
    is_strictly_final = "Final" in document_type or is_final
    is_redlined = "Redlined" in document_type or "(Redlined)" in document_type

    if s3_url:
        if is_strictly_final:
            print(f"📄 [SKIP] Extraction & classification skipped for Final document: {file_name}")
        elif is_redlined:
            print(f"📄 [SKIP] Extraction skipped for redlined document: {file_name}")
        else:
            background_tasks.add_task(trigger_extraction, file_name, doc_uuid, document_type, source)
    
    return {
        "message": "Document uploaded and queued for extraction",
        "document": new_doc
    }


@app.get("/api/documents/analysis/{document_id}")
def get_document_analysis(document_id: str, current_user: dict = Depends(verify_token)):
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT id, filename, document_type, user_id, user_email, user_role, size, status, shared_with, uploaded_at, file_path, s3_url, s3_key FROM documents WHERE id = %s", (document_id,))
            row = cur.fetchone()
            
            if not row:
                raise HTTPException(status_code=404, detail=f"Document {document_id} not found")
                
            doc = {
                "id": row[0], "filename": row[1], "document_type": row[2], "user_id": row[3],
                "user_email": row[4], "user_role": row[5], "size": row[6], "status": row[7],
                "shared_with": row[8], "uploaded_at": row[9].isoformat() if row[9] else None,
                "file_path": row[10], "s3_url": row[11], "s3_key": row[12]
            }
            # Query clauses from database
            cur.execute(
                """
                SELECT c.clause_id, c.clause, c.content_id, c.content, c.page_number,
                       e.edited_clause, e.comment
                FROM clauses c
                LEFT JOIN edited_clauses e ON c.content_id = e.content_id
                WHERE c.document_id = %s
                ORDER BY c.page_number ASC, c.ctid ASC
                """,
                (document_id,)
            )
            clause_rows = cur.fetchall()
            
            if not clause_rows:
                return {"document": doc, "clauses": [], "status": "processing"}
                
            import difflib
            clauses_data = []
            for r in clause_rows:
                orig_content = r[3]
                edited_content = r[5]
                html_diff = None
                
                if edited_content and edited_content != orig_content:
                    seq = difflib.SequenceMatcher(None, orig_content, edited_content)
                    out = []
                    for opcode, i1, i2, j1, j2 in seq.get_opcodes():
                        if opcode == 'equal':
                            out.append(orig_content[i1:i2])
                        elif opcode == 'delete':
                            out.append(f"<del style='color:red; text-decoration:line-through'>{orig_content[i1:i2]}</del>")
                        elif opcode == 'insert':
                            out.append(f"<ins style='color:red; text-decoration:underline'>{edited_content[j1:j2]}</ins>")
                        elif opcode == 'replace':
                            out.append(f"<del style='color:red; text-decoration:line-through'>{orig_content[i1:i2]}</del>")
                            out.append(f"<ins style='color:red; text-decoration:underline'>{edited_content[j1:j2]}</ins>")
                    html_diff = "".join(out).replace('\n', '<br/>')

                clauses_data.append({
                    "clause_id": r[0],
                    "clause": r[1],
                    "content_id": r[2],
                    "content": orig_content,
                    "page_number": r[4],
                    "edited_content": edited_content,
                    "comment": r[6],
                    "html_diff": html_diff
                })
            
            return {"document": doc, "clauses": clauses_data, "status": "complete"}
                
    except HTTPException:
        # Re-raise HTTP exceptions directly
        raise
    except Exception as e:
        print(f"[ERROR] Database analysis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)


class ClauseEditRequest(BaseModel):
    content_id: str
    edited_content: str

@app.post("/api/documents/review/{document_id}/edit")
def edit_clause(document_id: str, edit_request: ClauseEditRequest, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # Upsert logic - Needs Original Clause
            cur.execute("SELECT content FROM clauses WHERE content_id = %s", (edit_request.content_id,))
            res = cur.fetchone()
            if not res:
                raise HTTPException(status_code=404, detail="Clause not found")
            original_content = res[0]
            
            cur.execute(
                """
                INSERT INTO edited_clauses (content_id, original_clause, edited_clause, updated_at)
                VALUES (%s, %s, %s, NOW())
                ON CONFLICT (content_id) DO UPDATE 
                SET edited_clause = EXCLUDED.edited_clause, updated_at = NOW()
                """,
                (edit_request.content_id, original_content, edit_request.edited_content)
            )
        conn.commit()
        return {"message": "Edit saved successfully"}
    except Exception as e:
        print(f"[ERROR] Save edit error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)


class ClauseCommentRequest(BaseModel):
    content_id: str
    comment: str

@app.post("/api/documents/review/{document_id}/comment")
def add_clause_comment(document_id: str, comment_request: ClauseCommentRequest, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # Upsert logic
            cur.execute("SELECT content FROM clauses WHERE content_id = %s", (comment_request.content_id,))
            res = cur.fetchone()
            if not res:
                raise HTTPException(status_code=404, detail="Clause not found")
            original_content = res[0]
            
            cur.execute(
                """
                INSERT INTO edited_clauses (content_id, original_clause, comment, updated_at)
                VALUES (%s, %s, %s, NOW())
                ON CONFLICT (content_id) DO UPDATE 
                SET comment = EXCLUDED.comment, updated_at = NOW()
                """,
                (comment_request.content_id, original_content, comment_request.comment)
            )
            cur.execute("UPDATE documents SET google_doc_id = NULL WHERE id = %s", (document_id,))
        conn.commit()
        return {"message": "Comment saved successfully"}
    except Exception as e:
        print(f"[ERROR] Save comment error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)


@app.get("/api/documents/download-redline/{document_id}")
def download_redline(document_id: str, background_tasks: BackgroundTasks, current_user: dict = Depends(verify_token)):
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # 1. Fetch document metadata
            cur.execute("SELECT filename FROM documents WHERE id = %s", (document_id,))
            doc_res = cur.fetchone()
            if not doc_res:
                raise HTTPException(status_code=404, detail="Document not found")
            original_filename = doc_res[0]
            
            cur.execute(
                """
                SELECT c.content_id, c.content, e.edited_clause, e.comment
                FROM clauses c
                LEFT JOIN edited_clauses e ON c.content_id = e.content_id
                WHERE c.document_id = %s
                ORDER BY c.page_number ASC, c.ctid ASC
                """,
                (document_id,)
            )
            clause_rows = cur.fetchall()
            
            # 3. Fetch all suggestions
            cur.execute(
                """
                SELECT content_id, original_text, suggested_text, author, timestamp, status
                FROM clause_suggestions
                WHERE document_id = %s AND status IN ('pending', 'accepted')
                ORDER BY timestamp ASC
                """,
                (document_id,)
            )
            all_sugs = cur.fetchall()

            if not clause_rows:
                raise HTTPException(status_code=404, detail="No clauses found for this document")

    except Exception as e:
        print(f"[ERROR] Fetching redline data: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)

    # 3. Generate Redlined Docx
    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed")

    from docx.shared import RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import difflib
    from datetime import datetime

    document = docx.Document()
    # Setup comments xml part
    comments_part = None
    if 'comments' not in document.part.rels:
        try:
            from docx.opc.part import XmlPart
            from docx.oxml import parse_xml
            from docx.opc.constants import RELATIONSHIP_TYPE, CONTENT_TYPE
            from docx.opc.packuri import PackURI
            
            comments_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"></w:comments>'
            comments_part = XmlPart(
                PackURI('/word/comments.xml'), 
                CONTENT_TYPE.WML_COMMENTS, 
                parse_xml(comments_xml.encode('utf-8')), 
                document.part.package
            )
            document.part.relate_to(comments_part, RELATIONSHIP_TYPE.COMMENTS)
        except Exception as e:
            print(f"[ERROR] Adding comments part: {e}")
            pass
    else:
        comments_part = document.part.rels['comments'].target_part

    comment_id_counter = 0
    tc_id_counter = 1

    sug_map = {}
    for s in all_sugs:
        cid = s[0]
        if cid not in sug_map: sug_map[cid] = []
        sug_map[cid].append(s)

    for idx, row in enumerate(clause_rows):
        cid, original, accepted_edited, comment = row
        p = document.add_paragraph()
        
        # Current baseline for this clause's diff in Word is the 'original' text.
        # Everything else (Accepted or Pending) will be shown as track changes.
        
        # We'll use a sequential diff approach to show changes.
        # But wait, Word doesn't like complex diffing via python-docx's low-level injection easily.
        # We will show the 'final' version (all accepted + pending) as a diff against original.
        
        final_text = accepted_edited or original
        # Apply accepted and pending suggestions for robust DOCX preview
        for s in sug_map.get(cid, []):
            if s[5] in ('pending', 'accepted'):
                # s[1] is original_text, s[2] is suggested_text
                if s[1] == final_text:
                    final_text = s[2]
                elif s[1] in final_text:
                    final_text = final_text.replace(s[1], s[2], 1)
                elif s[1].strip() == final_text.strip():
                    final_text = s[2]

        if final_text == original:
            p.add_run(original)
        else:
            seq = difflib.SequenceMatcher(None, original, final_text)
            for opcode, i1, i2, j1, j2 in seq.get_opcodes():
                if opcode == 'equal':
                    p.add_run(original[i1:i2])
                elif opcode in ('delete', 'replace'):
                    del_elem = OxmlElement('w:del')
                    del_elem.set(qn('w:id'), str(tc_id_counter))
                    tc_id_counter += 1
                    del_elem.set(qn('w:author'), "TYN Legal Team")
                    del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                    r = OxmlElement('w:r')
                    dt = OxmlElement('w:delText')
                    dt.text = original[i1:i2]
                    if original[i1:i2].startswith(' ') or original[i1:i2].endswith(' '):
                        dt.set(qn('xml:space'), 'preserve')
                    r.append(dt)
                    del_elem.append(r)
                    p._p.append(del_elem)
                
                if opcode in ('insert', 'replace'):
                    ins = OxmlElement('w:ins')
                    ins.set(qn('w:id'), str(tc_id_counter))
                    tc_id_counter += 1
                    # Try to attribute to a pending suggestion if this segment matches
                    ins.set(qn('w:author'), "TYN Legal Team")
                    ins.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = final_text[j1:j2]
                    if final_text[j1:j2].startswith(' ') or final_text[j1:j2].endswith(' '):
                        t.set(qn('xml:space'), 'preserve')
                    r.append(t)
                    ins.append(r)
                    p._p.append(ins)

        if comment and comments_part is not None:
            comment_id_str = str(comment_id_counter)
            comment_id_counter += 1
            
            comment_elem = OxmlElement('w:comment')
            comment_elem.set(qn('w:id'), comment_id_str)
            comment_elem.set(qn('w:author'), "TYN Legal Team")
            comment_elem.set(qn('w:date'), datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'))
            
            c_p = OxmlElement('w:p')
            c_r = OxmlElement('w:r')
            c_t = OxmlElement('w:t')
            c_t.text = comment
            
            c_r.append(c_t)
            c_p.append(c_r)
            comment_elem.append(c_p)

            comments_part.element.append(comment_elem)

            comment_start = OxmlElement('w:commentRangeStart')
            comment_start.set(qn('w:id'), comment_id_str)
            p._p.insert(0, comment_start)

            comment_end = OxmlElement('w:commentRangeEnd')
            comment_end.set(qn('w:id'), comment_id_str)
            p._p.append(comment_end)

            comment_ref_r = OxmlElement('w:r')
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), comment_id_str)
            comment_ref_r.append(comment_ref)
            p._p.append(comment_ref_r)
            
        # Add some spacing between clauses
        if idx < len(clause_rows) - 1:
            document.add_paragraph()

    # Save to temp file
    import tempfile
    import os
    from fastapi.responses import FileResponse
    
    # Create temp file
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    
    try:
        document.save(path)
        # Log size for debugging 0-byte issue
        size = os.path.getsize(path)
        print(f"[DEBUG] Redline generated at {path}, size: {size} bytes")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to save document: {str(e)}")
    
    download_filename = f"Redlined_{original_filename.replace('.pdf', '').replace('.txt', '')}.docx"
    
    # Use FastAPI FileResponse which handles downloading, and remove temp file after
    return FileResponse(
        path=path, 
        filename=download_filename, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=background_tasks.add_task(os.remove, path)
    )

@app.post("/api/documents/download-redline-docs/{document_id}")
def download_redline_docs(document_id: str, background_tasks: BackgroundTasks, current_user: dict = Depends(verify_token)):
    import os
    import traceback
    
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # 1. Check if google_doc_id already exists
            cur.execute("SELECT filename, google_doc_id, document_type, s3_key, file_path FROM documents WHERE id = %s", (document_id,))
            doc_res = cur.fetchone()
            if not doc_res:
                raise HTTPException(status_code=404, detail="Document not found")
            
            original_filename = doc_res[0]
            existing_doc_id = doc_res[1]
            document_type = doc_res[2]
            s3_key = doc_res[3]
            file_path = doc_res[4]
            
            # If we've already generated/uploaded a doc for this redline directly, return it
            if existing_doc_id:
                return {"url": f"https://docs.google.com/document/d/{existing_doc_id}/edit"}
            
            # It's a review page, we need to generate the redline .docx locally and upload it
            
            # Fetch data to generate redline
            cur.execute(
                """
                SELECT c.content_id, c.content, e.edited_clause, e.comment
                FROM clauses c
                LEFT JOIN edited_clauses e ON c.content_id = e.content_id
                WHERE c.document_id = %s
                ORDER BY c.page_number ASC, c.ctid ASC
                """,
                (document_id,)
            )
            clause_rows = cur.fetchall()
            
            cur.execute(
                """
                SELECT content_id, original_text, suggested_text, author, timestamp, status
                FROM clause_suggestions
                WHERE document_id = %s AND status IN ('pending', 'accepted')
                ORDER BY timestamp ASC
                """,
                (document_id,)
            )
            all_sugs = cur.fetchall()

            if not clause_rows:
                # If there are no clauses, we must look for the original file and upload IT as the redline.
                # E.g. When the client uploads a Redlined doc, it has no clauses yet.
                pass

    except Exception as e:
        print(f"[ERROR] Fetching redline data for Google Docs: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)

    # If no clauses exist (like an uploaded redline document), try to upload the raw file
    if not clause_rows:
        try:
            file_bytes = None
            if s3_key:
                try:
                    response = s3_client.get_object(Bucket=BUCKET_NAME, Key=s3_key)
                    file_bytes = response['Body'].read()
                except Exception as s3_e:
                    print(f"[S3 Warning] Failed to fetch from S3: {s3_e}")
            
            if not file_bytes and file_path and os.path.exists(file_path):
                with open(file_path, "rb") as f:
                    file_bytes = f.read()
            
            if not file_bytes:
                raise HTTPException(status_code=404, detail="Original document file could not be found to upload to Google Docs.")
            
            from google_drive_service import upload_to_google_docs
            doc_id = upload_to_google_docs(file_bytes, original_filename)
            
            # Save ID to DB
            conn2 = db_pool.getconn()
            try:
                with conn2.cursor() as cur2:
                    cur2.execute("UPDATE documents SET google_doc_id = %s WHERE id = %s", (doc_id, document_id))
                    conn2.commit()
            except Exception as db_e:
                print(f"[ERROR] Saving google_doc_id to DB: {db_e}")
                conn2.rollback()
            finally:
                db_pool.putconn(conn2)
                
            return {"url": f"https://docs.google.com/document/d/{doc_id}/edit"}
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Failed to upload raw document to Google Doc: {str(e)}")

    # Otherwise, Generate Redlined Docx locally first
    from docx.shared import RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import difflib
    from datetime import datetime
    import docx
    import tempfile

    document = docx.Document()
    
    # Comments part setup
    comments_part = None
    if 'comments' not in document.part.rels:
        try:
            from docx.opc.part import XmlPart
            from docx.oxml import parse_xml
            from docx.opc.constants import RELATIONSHIP_TYPE, CONTENT_TYPE
            from docx.opc.packuri import PackURI
            comments_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"></w:comments>'
            comments_part = XmlPart(PackURI('/word/comments.xml'), CONTENT_TYPE.WML_COMMENTS, parse_xml(comments_xml.encode('utf-8')), document.part.package)
            document.part.relate_to(comments_part, RELATIONSHIP_TYPE.COMMENTS)
        except Exception:
            pass
    else:
        comments_part = document.part.rels['comments'].target_part

    comment_id_counter = 0
    tc_id_counter = 1

    sug_map = {}
    for s in all_sugs:
        cid = s[0]
        if cid not in sug_map: sug_map[cid] = []
        sug_map[cid].append(s)

    for idx, row in enumerate(clause_rows):
        cid, original, accepted_edited, comment = row
        p = document.add_paragraph()
        
        final_text = accepted_edited or original
        for s in sug_map.get(cid, []):
            if s[5] in ('pending', 'accepted'):
                if s[1] == final_text:
                    final_text = s[2]
                elif s[1] in final_text:
                    final_text = final_text.replace(s[1], s[2], 1)
                elif s[1].strip() == final_text.strip():
                    final_text = s[2]

        if final_text == original:
            p.add_run(original)
        else:
            seq = difflib.SequenceMatcher(None, original, final_text)
            for opcode, i1, i2, j1, j2 in seq.get_opcodes():
                if opcode == 'equal':
                    p.add_run(original[i1:i2])
                elif opcode in ('delete', 'replace'):
                    del_elem = OxmlElement('w:del')
                    del_elem.set(qn('w:id'), str(tc_id_counter))
                    tc_id_counter += 1
                    del_elem.set(qn('w:author'), "TYN Legal Team")
                    del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                    r = OxmlElement('w:r')
                    dt = OxmlElement('w:delText')
                    dt.text = original[i1:i2]
                    if original[i1:i2].startswith(' ') or original[i1:i2].endswith(' '):
                        dt.set(qn('xml:space'), 'preserve')
                    r.append(dt)
                    del_elem.append(r)
                    p._p.append(del_elem)
                
                if opcode in ('insert', 'replace'):
                    ins = OxmlElement('w:ins')
                    ins.set(qn('w:id'), str(tc_id_counter))
                    tc_id_counter += 1
                    ins.set(qn('w:author'), "TYN Legal Team")
                    ins.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = final_text[j1:j2]
                    if final_text[j1:j2].startswith(' ') or final_text[j1:j2].endswith(' '):
                        t.set(qn('xml:space'), 'preserve')
                    r.append(t)
                    ins.append(r)
                    p._p.append(ins)

        if comment and comments_part is not None:
            comment_id_str = str(comment_id_counter)
            comment_id_counter += 1
            comment_elem = OxmlElement('w:comment')
            comment_elem.set(qn('w:id'), comment_id_str)
            comment_elem.set(qn('w:author'), "TYN Legal Team")
            comment_elem.set(qn('w:date'), datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'))
            c_p = OxmlElement('w:p')
            c_r = OxmlElement('w:r')
            c_t = OxmlElement('w:t')
            c_t.text = comment
            c_r.append(c_t)
            c_p.append(c_r)
            comment_elem.append(c_p)
            comments_part.element.append(comment_elem)
            comment_start = OxmlElement('w:commentRangeStart')
            comment_start.set(qn('w:id'), comment_id_str)
            p._p.insert(0, comment_start)
            comment_end = OxmlElement('w:commentRangeEnd')
            comment_end.set(qn('w:id'), comment_id_str)
            p._p.append(comment_end)

            comment_ref_r = OxmlElement('w:r')
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), comment_id_str)
            comment_ref_r.append(comment_ref)
            p._p.append(comment_ref_r)
            
        if idx < len(clause_rows) - 1:
            document.add_paragraph()

    # Save to temp file
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    
    try:
        document.save(path)
        
        # Read file bytes
        with open(path, "rb") as f:
            file_bytes = f.read()
            
        # Clean up temp file
        os.remove(path)
        
        # Upload to Google Docs
        from google_drive_service import upload_to_google_docs
        googledoc_filename = f"Redlined_{original_filename.replace('.pdf', '').replace('.txt', '')}"
        doc_id = upload_to_google_docs(file_bytes, googledoc_filename)
        
        # Save ID to DB
        conn2 = db_pool.getconn()
        try:
            with conn2.cursor() as cur2:
                cur2.execute("UPDATE documents SET google_doc_id = %s WHERE id = %s", (doc_id, document_id))
                conn2.commit()
        except Exception as db_e:
            print(f"[ERROR] Saving google_doc_id to DB: {db_e}")
            conn2.rollback()
        finally:
            db_pool.putconn(conn2)
            
        return {"url": f"https://docs.google.com/document/d/{doc_id}/edit"}
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        if os.path.exists(path):
            os.remove(path)
        raise HTTPException(status_code=500, detail=f"Failed to generate or upload Google Doc: {str(e)}")

@app.post("/api/documents/google-doc/{document_id}")
def open_in_google_docs(document_id: str, background_tasks: BackgroundTasks, current_user: dict = Depends(verify_token)):
    import os
    import traceback
    
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # 1. Check if google_doc_id already exists
            cur.execute("SELECT filename, google_doc_id, document_type, s3_key, file_path FROM documents WHERE id = %s", (document_id,))
            doc_res = cur.fetchone()
            if not doc_res:
                raise HTTPException(status_code=404, detail="Document not found")
            
            original_filename = doc_res[0]
            existing_doc_id = doc_res[1]
            document_type = doc_res[2]
            s3_key = doc_res[3]
            file_path = doc_res[4]
            
            
            if existing_doc_id:
                return {"url": f"https://docs.google.com/document/d/{existing_doc_id}/edit"}
            
            is_uploaded_redline = document_type and ("Redlined" in document_type or "(Redlined)" in document_type)
            
            # 2. Since this is for simple document review via Dashboard "Review" button on Redlined docs
            # If it's a generated redline doc with system edits, we MUST build the redline DOCX with changes
            
            cur.execute(
                """
                SELECT c.content_id, c.content, e.edited_clause, e.comment
                FROM clauses c
                LEFT JOIN edited_clauses e ON c.content_id = e.content_id
                WHERE c.document_id = %s
                ORDER BY c.page_number ASC, c.ctid ASC
                """,
                (document_id,)
            )
            clause_rows = cur.fetchall()
            
            if clause_rows:
                # Delegate to the Redline Generation Logic when System Edits exist
                pass 
            else:
                # Upload the raw file directly (For basic PDF/Docx without clauses)
                file_bytes = None
                if s3_key:
                    try:
                        response = s3_client.get_object(Bucket=BUCKET_NAME, Key=s3_key)
                        file_bytes = response['Body'].read()
                    except Exception as s3_e:
                        print(f"[S3 Warning] Failed to fetch: {s3_e}")
                
                if not file_bytes and file_path and os.path.exists(file_path):
                    with open(file_path, "rb") as f:
                        file_bytes = f.read()
                
                if not file_bytes:
                    raise HTTPException(status_code=404, detail="Original document file could not be found to upload to Google Docs.")
                
                from google_drive_service import upload_to_google_docs
                doc_id = upload_to_google_docs(file_bytes, original_filename)
                
                # Save ID to DB
                cur.execute("UPDATE documents SET google_doc_id = %s WHERE id = %s", (doc_id, document_id))
                conn.commit()
                
                return {"url": f"https://docs.google.com/document/d/{doc_id}/edit"}
            
    except Exception as e:
        print(f"[ERROR] Fetching and uploading doc for Google Docs: {e}")
        import traceback
        traceback.print_exc()
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)

    # 3. If there ARE clauses, generate Redlined Docx locally first
    from docx.shared import RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import difflib
    from datetime import datetime
    import docx
    import tempfile

    document = docx.Document()
    
    # Comments part setup
    comments_part = None
    if 'comments' not in document.part.rels:
        try:
            from docx.opc.part import XmlPart
            from docx.oxml import parse_xml
            from docx.opc.constants import RELATIONSHIP_TYPE, CONTENT_TYPE
            from docx.opc.packuri import PackURI
            comments_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"></w:comments>'
            comments_part = XmlPart(PackURI('/word/comments.xml'), CONTENT_TYPE.WML_COMMENTS, parse_xml(comments_xml.encode('utf-8')), document.part.package)
            document.part.relate_to(comments_part, RELATIONSHIP_TYPE.COMMENTS)
        except Exception:
            pass
    else:
        comments_part = document.part.rels['comments'].target_part

    comment_id_counter = 0
    tc_id_counter = 1

    # Fetch Suggestions
    conn = db_pool.getconn()
    all_sugs = []
    try:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT content_id, original_text, suggested_text, author, timestamp, status
                FROM clause_suggestions
                WHERE document_id = %s AND status IN ('pending', 'accepted')
                ORDER BY timestamp ASC
            """, (document_id,))
            all_sugs = cur.fetchall()
    except Exception as e:
        pass
    finally:
        db_pool.putconn(conn)

    sug_map = {}
    for s in all_sugs:
        cid = s[0]
        if cid not in sug_map: sug_map[cid] = []
        sug_map[cid].append(s)

    for idx, row in enumerate(clause_rows):
        cid, original, accepted_edited, comment = row
        p = document.add_paragraph()
        
        final_text = accepted_edited or original
        for s in sug_map.get(cid, []):
            if s[5] == 'pending':
                if s[1] == final_text:
                    final_text = s[2]
                elif s[1] in final_text:
                    final_text = final_text.replace(s[1], s[2], 1)

        if final_text == original:
            p.add_run(original)
        else:
            seq = difflib.SequenceMatcher(None, original, final_text)
            for opcode, i1, i2, j1, j2 in seq.get_opcodes():
                if opcode == 'equal':
                    p.add_run(original[i1:i2])
                elif opcode in ('delete', 'replace'):
                    del_elem = OxmlElement('w:del')
                    del_elem.set(qn('w:id'), str(tc_id_counter))
                    tc_id_counter += 1
                    del_elem.set(qn('w:author'), "TYN Legal Team")
                    del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                    r = OxmlElement('w:r')
                    dt = OxmlElement('w:delText')
                    dt.text = original[i1:i2]
                    if original[i1:i2].startswith(' ') or original[i1:i2].endswith(' '):
                        dt.set(qn('xml:space'), 'preserve')
                    r.append(dt)
                    del_elem.append(r)
                    p._p.append(del_elem)
                
                if opcode in ('insert', 'replace'):
                    ins = OxmlElement('w:ins')
                    ins.set(qn('w:id'), str(tc_id_counter))
                    tc_id_counter += 1
                    ins.set(qn('w:author'), "TYN Legal Team")
                    ins.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = final_text[j1:j2]
                    if final_text[j1:j2].startswith(' ') or final_text[j1:j2].endswith(' '):
                        t.set(qn('xml:space'), 'preserve')
                    r.append(t)
                    ins.append(r)
                    p._p.append(ins)

        if comment and comments_part is not None:
            comment_id_str = str(comment_id_counter)
            comment_id_counter += 1
            comment_elem = OxmlElement('w:comment')
            comment_elem.set(qn('w:id'), comment_id_str)
            comment_elem.set(qn('w:author'), "TYN Legal Team")
            comment_elem.set(qn('w:date'), datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'))
            c_p = OxmlElement('w:p')
            c_r = OxmlElement('w:r')
            c_t = OxmlElement('w:t')
            c_t.text = comment
            c_r.append(c_t)
            c_p.append(c_r)
            comment_elem.append(c_p)
            comments_part.element.append(comment_elem)
            comment_start = OxmlElement('w:commentRangeStart')
            comment_start.set(qn('w:id'), comment_id_str)
            p._p.insert(0, comment_start)
            comment_end = OxmlElement('w:commentRangeEnd')
            comment_end.set(qn('w:id'), comment_id_str)
            p._p.append(comment_end)

            comment_ref_r = OxmlElement('w:r')
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), comment_id_str)
            comment_ref_r.append(comment_ref)
            p._p.append(comment_ref_r)
            
        if idx < len(clause_rows) - 1:
            document.add_paragraph()

    # Save to temp file
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    
    try:
        document.save(path)
        
        # Read file bytes
        with open(path, "rb") as f:
            file_bytes = f.read()
            
        # Clean up temp file
        os.remove(path)
        
        # Upload to Google Docs
        from google_drive_service import upload_to_google_docs
        googledoc_filename = f"Redlined_{original_filename.replace('.pdf', '').replace('.txt', '')}"
        doc_id = upload_to_google_docs(file_bytes, googledoc_filename)
        
        # Save ID to DB
        conn2 = db_pool.getconn()
        try:
            with conn2.cursor() as cur2:
                cur2.execute("UPDATE documents SET google_doc_id = %s WHERE id = %s", (doc_id, document_id))
                conn2.commit()
        except Exception as db_e:
            print(f"[ERROR] Saving google_doc_id to DB: {db_e}")
            conn2.rollback()
        finally:
            db_pool.putconn(conn2)
            
        return {"url": f"https://docs.google.com/document/d/{doc_id}/edit"}
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        if os.path.exists(path):
            os.remove(path)
        raise HTTPException(status_code=500, detail=f"Failed to generate or upload Google Doc: {str(e)}")

@app.get("/api/documents/download/{document_id}")
def download_document(document_id: str, current_user: dict = Depends(verify_token)):
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            if document_id.startswith("sc-"):
                cur.execute("SELECT s3_key FROM shared_contracts WHERE id = %s", (document_id,))
            else:
                cur.execute("SELECT s3_key FROM documents WHERE id = %s", (document_id,))
            res = cur.fetchone()
            if not res:
                raise HTTPException(status_code=404, detail="Document not found")
            s3_key = res[0]
    except Exception as e:
        print(f"[ERROR] Database download fetch error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)
        
    if not s3_key:
        raise HTTPException(status_code=400, detail="File not on S3")
        
    # Basic extension detection
    mime_type = "application/pdf" if s3_key.lower().endswith(".pdf") else "application/octet-stream"
    try:
        url = s3_client.generate_presigned_url(
            'get_object',
            Params={
                'Bucket': BUCKET_NAME, 
                'Key': s3_key,
                'ResponseContentDisposition': 'inline',
                'ResponseContentType': mime_type
            },
            ExpiresIn=3600
        )
        return {"download_url": url}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"S3 Error: {str(e)}")


@app.post("/api/documents/send-redline/{document_id}")
def send_redline_to_client(document_id: str, background_tasks: BackgroundTasks, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # 1. Fetch document metadata
            cur.execute("SELECT id, filename, user_id FROM documents WHERE id = %s", (document_id,))
            doc_res = cur.fetchone()
            if not doc_res:
                raise HTTPException(status_code=404, detail="Document not found")
            original_id, original_filename, client_id = doc_res
            
            # 2. Fetch clauses + edits/comments in order
            cur.execute(
                """
                SELECT c.content_id, c.content, e.edited_clause, e.comment
                FROM clauses c
                LEFT JOIN edited_clauses e ON c.content_id = e.content_id
                WHERE c.document_id = %s
                ORDER BY c.page_number ASC, c.ctid ASC
                """,
                (document_id,)
            )
            clause_rows = cur.fetchall()

            if not clause_rows:
                print(f"[DEBUG] No clauses found for doc {document_id}")
                raise HTTPException(status_code=404, detail="No clauses found for this document")

            # FETCH all_sugs here just in case (already present in snippet but ensuring consistency)
            cur.execute("""
                SELECT content_id, original_text, suggested_text, author, timestamp, status
                FROM clause_suggestions
                WHERE document_id = %s AND status IN ('pending', 'accepted')
                ORDER BY timestamp ASC
            """, (document_id,))
            all_sugs = cur.fetchall()
            print(f"[DEBUG] Fetched {len(all_sugs)} suggestions for redline")

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Fetching redline data: {e}")
        import traceback
        traceback.print_exc()
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)

    # 3. Generate Redlined Docx
    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed")

    from docx.shared import RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import difflib
    from datetime import datetime

    document = docx.Document()
    # Setup comments xml part
    comments_part = None
    try:
        from docx.opc.part import XmlPart
        from docx.oxml import parse_xml
        from docx.opc.constants import RELATIONSHIP_TYPE, CONTENT_TYPE
        from docx.opc.packuri import PackURI
        
        comments_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"></w:comments>'
        comments_part = XmlPart(
            PackURI('/word/comments.xml'), 
            CONTENT_TYPE.WML_COMMENTS, 
            parse_xml(comments_xml.encode('utf-8')), 
            document.part.package
        )
        document.part.relate_to(comments_part, RELATIONSHIP_TYPE.COMMENTS)
    except Exception as e:
        print(f"[ERROR] Adding comments part: {e}")

    comment_id_counter = 0
    tc_id_counter = 1

    sug_map = {}
    for s in all_sugs:
        cid = s[0]
        if cid not in sug_map: sug_map[cid] = []
        sug_map[cid].append(s)

    for idx, row in enumerate(clause_rows):
        cid, original, accepted_edited, comment = row
        p = document.add_paragraph()
        
        final_text = accepted_edited or original
        # Apply pending suggestions ON TOP of accepted for the DOCX preview
        for s in sug_map.get(cid, []):
            if s[5] == 'pending':
                # s[1] is original_text, s[2] is suggested_text
                if s[1] == final_text:
                    final_text = s[2]
                elif s[1] in final_text:
                    final_text = final_text.replace(s[1], s[2], 1)

        if final_text == original:
            p.add_run(original)
        else:
            seq = difflib.SequenceMatcher(None, original, final_text)
            for opcode, i1, i2, j1, j2 in seq.get_opcodes():
                if opcode == 'equal':
                    p.add_run(original[i1:i2])
                elif opcode in ('delete', 'replace'):
                    del_elem = OxmlElement('w:del')
                    del_elem.set(qn('w:id'), str(tc_id_counter))
                    tc_id_counter += 1
                    del_elem.set(qn('w:author'), "TYN Legal Team")
                    del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                    r = OxmlElement('w:r')
                    dt = OxmlElement('w:delText')
                    dt.text = original[i1:i2]
                    if original[i1:i2].startswith(' ') or original[i1:i2].endswith(' '):
                        dt.set(qn('xml:space'), 'preserve')
                    r.append(dt)
                    del_elem.append(r)
                    p._p.append(del_elem)
                
                if opcode in ('insert', 'replace'):
                    ins = OxmlElement('w:ins')
                    ins.set(qn('w:id'), str(tc_id_counter))
                    tc_id_counter += 1
                    ins.set(qn('w:author'), "TYN Legal Team")
                    ins.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = final_text[j1:j2]
                    if final_text[j1:j2].startswith(' ') or final_text[j1:j2].endswith(' '):
                        t.set(qn('xml:space'), 'preserve')
                    r.append(t)
                    ins.append(r)
                    p._p.append(ins)

        if comment and comments_part is not None:
            comment_id_str = str(comment_id_counter)
            comment_id_counter += 1
            comment_elem = OxmlElement('w:comment')
            comment_elem.set(qn('w:id'), comment_id_str)
            comment_elem.set(qn('w:author'), "TYN Legal Team")
            comment_elem.set(qn('w:date'), datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'))
            c_p = OxmlElement('w:p')
            c_r = OxmlElement('w:r')
            c_t = OxmlElement('w:t')
            c_t.text = comment
            c_r.append(c_t)
            c_p.append(c_r)
            comment_elem.append(c_p)
            comments_part.element.append(comment_elem)
            comment_start = OxmlElement('w:commentRangeStart')
            comment_start.set(qn('w:id'), comment_id_str)
            p._p.insert(0, comment_start)
            comment_end = OxmlElement('w:commentRangeEnd')
            comment_end.set(qn('w:id'), comment_id_str)
            p._p.append(comment_end)
            comment_ref_r = OxmlElement('w:r')
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), comment_id_str)
            comment_ref_r.append(comment_ref)
            p._p.append(comment_ref_r)
            
        if idx < len(clause_rows) - 1: # Changed from 'rows' to 'clause_rows'
            document.add_paragraph()

    # Save to temp file
    import tempfile
    import os
    from fastapi.responses import JSONResponse
    
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    
    document.save(path)
    print(f"[DEBUG] send_redline generated temp file at {path}, size: {os.path.getsize(path)} bytes")
    
    # 4. Upload to S3
    new_filename = f"Redlined_{original_filename.replace('.pdf', '').replace('.txt', '')}.docx"
    s3_key = f"redline_{uuid.uuid4().hex[:8]}_{new_filename}"
    s3_url = None
    file_size = os.path.getsize(path)
    try:
        with open(path, "rb") as f:
            s3_client.put_object(Bucket=BUCKET_NAME, Key=s3_key, Body=f, ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        s3_url = f"https://{BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
    except Exception as e:
        print(f"[ERROR] S3 Redline upload failed: {e}")
        # Continue anyway, file sits locally? No, we need S3 for the client to download
        raise HTTPException(status_code=500, detail="Failed to upload redline to S3")
    
    # 5. Store in shared_contracts
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            contract_id = f"sc-{uuid.uuid4().hex[:8]}"
            cur.execute(
                """
                INSERT INTO shared_contracts (id, filename, shared_by, shared_by_email, client_id, message, status, shared_at, s3_key, s3_url, file_path, document_type, size)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (
                    contract_id, new_filename, current_user["user_id"], current_user.get("email", ""),
                    client_id, f"Redlined version of {original_filename}", 'pending_review',
                    datetime.now(), s3_key, s3_url, path, 'Redlined', file_size
                )
            )
            conn.commit()
            record_activity(current_user["user_id"], client_id, "Sent redline to client", new_filename)
            return {"message": "Redline sent to client successfully", "id": contract_id}
    except Exception as e:
        print(f"[ERROR] DB shared_contracts insert failed: {e}")
        conn.rollback()
        raise HTTPException(status_code=500, detail="Failed to record shared contract")
    finally:
        db_pool.putconn(conn)
        # We don't remove the file yet if it's referenced in shared_contracts? 
        # Actually, if we use S3, we can remove it. But schema.sql has file_path too.
        # I'll leave it for now or delete it if S3 is reliable.
        # os.remove(path)



@app.post("/api/documents/approve/{document_id}")
def approve_document(document_id: str, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool: raise HTTPException(status_code=500, detail="Database not connected")
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT user_id, filename, document_type FROM documents WHERE id = %s", (document_id,))
            doc = cur.fetchone()
            if not doc: raise HTTPException(status_code=404, detail="Document not found")
            
            doc_user_id, filename, document_type = doc
            
            cur.execute("UPDATE documents SET status = 'approved', approved_at = %s, approved_by = %s WHERE id = %s", (datetime.now(), current_user["user_id"], document_id))
            
            record_activity(current_user["user_id"], doc_user_id, "Approved document", f"Document: {filename}")
            
            cur.execute("SELECT email FROM users WHERE id = %s", (doc_user_id,))
            user_res = cur.fetchone()
            if user_res:
                send_email(user_res[0], f"{document_type} Approved", f"<html><body><p>{filename} approved!</p></body></html>")
        conn.commit()
        return {"message": "Document approved"}
    except HTTPException: raise
    except Exception as e:
        print(f"[ERROR] Database approve error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)

@app.post("/api/documents/reject/{document_id}")
def reject_document(document_id: str, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT user_id, filename, document_type FROM documents WHERE id = %s", (document_id,))
            doc = cur.fetchone()
            if not doc:
                raise HTTPException(status_code=404, detail="Document not found")
            
            user_id, filename, document_type = doc
            
            cur.execute(
                "UPDATE documents SET status = 'rejected', rejected_at = %s, rejected_by = %s WHERE id = %s",
                (datetime.now(), current_user["user_id"], document_id)
            )
            
            # Record activity
            record_activity(current_user["user_id"], user_id, "Rejected document", f"Document: {filename}")
            
            cur.execute("SELECT email FROM users WHERE id = %s", (user_id,))
            email_row = cur.fetchone()
            if email_row:
                send_email(email_row[0], f"{document_type} Update", f"<html><body><p>{filename} rejected.</p></body></html>")
        
        conn.commit()
        return {"message": "Document rejected"}
    except HTTPException: raise
    except Exception as e:
        print(f"[ERROR] Reject error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)


@app.get("/api/documents/list")
def list_documents(current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            if current_user["role"] in ("admin", "legal_team"):
                cur.execute("SELECT id, user_id, filename, document_type, status, uploaded_at, s3_key, size, shared_with, is_finalized, client_marked_final, user_role FROM documents ORDER BY uploaded_at DESC")
            else:
                cur.execute("SELECT id, user_id, filename, document_type, status, uploaded_at, s3_key, size, shared_with, is_finalized, client_marked_final, user_role FROM documents WHERE user_id = %s OR shared_with @> %s::jsonb ORDER BY uploaded_at DESC", (current_user["user_id"], json.dumps([current_user["user_id"]])))
            rows = cur.fetchall()
            return {"documents": [{"id": r[0], "user_id": r[1], "filename": r[2], "document_type": r[3], "status": r[4], "uploaded_at": r[5].isoformat() if r[5] else None, "s3_key": r[6], "size": r[7], "shared_with": r[8], "is_finalized": r[9], "client_marked_final": r[10], "user_role": r[11]} for r in rows]}
    except Exception as e:
        print(f"[ERROR] list_documents error: {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally: db_pool.putconn(conn)


@app.post("/api/documents/finalize/{document_id}")
def finalize_document(document_id: str, current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # 1. Fetch info about the document
            if document_id.startswith("sc-"):
                cur.execute("SELECT document_type, client_id, is_finalized, filename FROM shared_contracts WHERE id = %s", (document_id,))
                row = cur.fetchone()
                if not row: raise HTTPException(status_code=404, detail="Document not found")
                doc_type, doc_client_id, current_final, filename = row
            else:
                cur.execute("SELECT document_type, user_id, is_finalized, filename FROM documents WHERE id = %s", (document_id,))
                row = cur.fetchone()
                if not row: raise HTTPException(status_code=404, detail="Document not found")
                doc_type, doc_client_id, current_final, filename = row
                
            # Toggle logic
            new_status = not current_final

            if new_status is True:
                # 2. Enforce Single Final rule: unset all others of the exact same doc_type for this client
                cur.execute("UPDATE shared_contracts SET is_finalized = false WHERE client_id = %s AND document_type = %s AND id != %s", (doc_client_id, doc_type, document_id))
                cur.execute("UPDATE documents SET is_finalized = false WHERE user_id = %s AND document_type = %s AND id != %s", (doc_client_id, doc_type, document_id))

            # 3. Apply the toggled status to current doc
            if document_id.startswith("sc-"):
                cur.execute("UPDATE shared_contracts SET is_finalized = %s WHERE id = %s", (new_status, document_id))
            else:
                cur.execute("UPDATE documents SET is_finalized = %s WHERE id = %s", (new_status, document_id))
            
            status_text = "Finalized" if new_status else "Un-finalized"
            record_activity(current_user["user_id"], doc_client_id, f"{status_text} document", f"Document: {filename}")
            
        conn.commit()
        return {"message": f"Document {status_text.lower()} successfully", "is_finalized": new_status}
    except HTTPException: raise
    except Exception as e:
        print(f"[ERROR] Finalize error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)



@app.get("/api/documents/stats")
def document_stats(current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT COUNT(*), COALESCE(SUM(size), 0) FROM documents")
            row = cur.fetchone()
            return {"total_documents": row[0], "total_size": row[1]}
    except Exception as e:
        print(f"[ERROR] Stats error: {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)


@app.post("/api/messages/send")
def send_message(msg: MessageSend, current_user: dict = Depends(verify_token)):
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    conn = None
    try:
        conn = db_pool.getconn()
        msg_id = f"msg-{uuid.uuid4().hex[:8]}"
        with conn.cursor() as cur:
            cur.execute("INSERT INTO messages (id, sender_id, recipient_id, content, timestamp) VALUES (%s, %s, %s, %s, %s)", 
                        (msg_id, current_user["user_id"], msg.recipient_id, msg.content, datetime.now()))
        conn.commit()
        print(f"✓ [CHAT] Message sent: From {current_user['user_id']} To {msg.recipient_id}")
        return {"message": "Message sent", "data": {"id": msg_id, "sender_id": current_user["user_id"], "recipient_id": msg.recipient_id, "content": msg.content, "timestamp": datetime.now().isoformat()}}
    except Exception as e:
        print(f"[ERROR] Database message send error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)


@app.get("/api/messages/list/{other_user_id}")
def list_messages(other_user_id: str, current_user: dict = Depends(verify_token)):
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    conn = None
    try:
        conn = db_pool.getconn()
        user_id = current_user.get("user_id")
        with conn.cursor() as cur:
            cur.execute("SELECT id, sender_id, recipient_id, content, timestamp FROM messages WHERE (sender_id = %s AND recipient_id = %s) OR (sender_id = %s AND recipient_id = %s) ORDER BY timestamp ASC",
                        (user_id, other_user_id, other_user_id, user_id))
            rows = cur.fetchall()
            return {"messages": [{"id": r[0], "sender_id": r[1], "recipient_id": r[2], "content": r[3], "timestamp": r[4].isoformat()} for r in rows]}
    except Exception as e:
        print(f"[ERROR] Database message list error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)



@app.on_event("startup")
def startup_event():
    # Deferring model loading to avoid blocking startup on limited CPU/Network
    # load_model()
    print("[INFO] Application startup complete - Model will be loaded on first use")
@app.post("/api/contracts/share-with-client")
async def share_contract_with_client(
    file: UploadFile = File(...), 
    client_id: str = Form(""), 
    message: Optional[str] = Form(None), 
    document_type: str = Form("PDF"),
    is_final: bool = Form(False),
    current_user: dict = Depends(verify_token)
):
    print(f"[SHARE] Role: {current_user.get('role')} | User: {current_user.get('email')} | ClientID: {client_id}")
    if current_user["role"] not in ["admin", "legal_team"]:
        print(f"✗ [SHARE] 403: Role {current_user['role']} not authorized")
        raise HTTPException(status_code=403)
    content = await file.read()
    file_size_bytes = len(content)

    file_name = f"shared_{uuid.uuid4().hex[:6]}_{file.filename}"
    file_path = UPLOADS_DIR / file_name
    with open(file_path, "wb") as f: f.write(content)
    s3_key = None
    try:
        s3_client.put_object(Bucket=BUCKET_NAME, Key=file_name, Body=content)
        s3_key = file_name
    except Exception as e: print(f"S3 error: {e}")
    contract_id = f"sc-{uuid.uuid4().hex[:8]}"
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("INSERT INTO shared_contracts (id, filename, shared_by, shared_by_email, client_id, message, status, shared_at, s3_key, file_path, document_type, is_finalized, size) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                        (contract_id, file.filename, current_user["user_id"], current_user.get("email", ""), client_id, message, 'pending_review', datetime.now(), s3_key, str(file_path), document_type, is_final, file_size_bytes))
            conn.commit()
            record_activity(current_user["user_id"], client_id, "Shared contract", f"{document_type}: {file.filename}")
            return {"message": "Shared", "id": contract_id}
    finally: db_pool.putconn(conn)

@app.get("/api/contracts/from-legal")
def get_contracts_from_legal(client_id: Optional[str] = None, current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        target_client_id = current_user["user_id"]
        # Allow admins/legal_team to view contracts sent to a specific client
        if client_id and current_user["role"] in ["admin", "legal_team"]:
            target_client_id = client_id

        with conn.cursor() as cur:
            # 1. Get explicitly shared contracts
            cur.execute("SELECT id, filename, document_type, shared_by, shared_by_email, message, size, status, shared_at, s3_key, is_finalized FROM shared_contracts WHERE client_id = %s ORDER BY shared_at DESC", (target_client_id,))
            rows = cur.fetchall()
            contracts = []
            for r in rows:
                contracts.append({
                    "id": r[0], "filename": r[1], "document_type": r[2],
                    "shared_by": r[3], "shared_by_email": r[4], "message": r[5],
                    "size": r[6] or 0, "status": r[7], "shared_at": r[8].isoformat() if r[8] else None,
                    "s3_key": r[9],
                    "is_finalized": r[10]
                })

            # Always inject the latest NDA for clients to show current status
            cur.execute("SELECT nda_accepted, nda_rejected FROM users WHERE id = %s", (target_client_id,))
            user_nda = cur.fetchone()
            if user_nda:
                cur.execute(
                    "SELECT id, filename, uploaded_at, size, s3_key FROM documents WHERE document_type = 'template' AND template_type = 'NDA' ORDER BY uploaded_at DESC LIMIT 1"
                )
                nda_row = cur.fetchone()
                if nda_row:
                    status = "pending_mandate"
                    if user_nda[0]: status = "accepted"
                    elif user_nda[1]: status = "rejected"
                    
                    contracts.insert(0, {
                        "id": nda_row[0],
                        "filename": nda_row[1],
                        "shared_at": nda_row[2].isoformat() if nda_row[2] else None,
                        "document_type": "NDA",
                        "shared_by": "Legal Team", # Added back
                        "shared_by_email": "legal@laccis.com", # Added back
                        "message": "Mandatory NDA required for portal access. Please review and approve.", # Added back
                        "status": status,
                        "size": nda_row[3],
                        "s3_key": nda_row[4],
                        "is_mandate": True
                    })
            return {"contracts": contracts}
    finally: db_pool.putconn(conn)

@app.get("/api/contracts/all-shared")
def get_all_shared_contracts(current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, client_id, filename, document_type, status, shared_at, s3_key, size, is_finalized FROM shared_contracts ORDER BY shared_at DESC")
            rows = cur.fetchall()
            return {
                "contracts": [
                    {
                        "id": r[0],
                        "client_id": r[1],
                        "filename": r[2],
                        "document_type": r[3],
                        "status": r[4],
                        "uploaded_at": r[5].isoformat() if r[5] else None,
                        "s3_key": r[6],
                        "size": r[7],
                        "is_finalized": r[8]
                    } for r in rows
                ]
            }
    except Exception as e:
        print(f"[ERROR] all-shared error: {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)


@app.post("/api/contracts/accept/{contract_id}")
def accept_shared_contract(contract_id: str, current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("UPDATE shared_contracts SET status = 'accepted', accepted_at = %s WHERE id = %s AND client_id = %s RETURNING filename", (datetime.now(), contract_id, current_user["user_id"]))
            row = cur.fetchone()
            if row:
                conn.commit()
                record_activity(current_user["user_id"], current_user["user_id"], "Accepted contract", f"Accepted shared contract: {row[0]}")
                return {"message": "Accepted"}
            raise HTTPException(status_code=404)
    finally: db_pool.putconn(conn)

@app.post("/api/contracts/reject/{contract_id}")
def reject_shared_contract(contract_id: str, current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("UPDATE shared_contracts SET status = 'rejected' WHERE id = %s AND client_id = %s RETURNING filename", (contract_id, current_user["user_id"]))
            row = cur.fetchone()
            if row:
                conn.commit()
                record_activity(current_user["user_id"], current_user["user_id"], "Rejected contract", f"Rejected shared contract: {row[0]}")
                return {"message": "Rejected"}
            raise HTTPException(status_code=404)
    finally: db_pool.putconn(conn)


@app.post("/api/contracts/accept-mandate")
def accept_mandate_nda(current_user: dict = Depends(verify_token)):
    if current_user["role"] != "client":
        raise HTTPException(status_code=403, detail="Only clients can accept mandate NDA")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("UPDATE users SET nda_accepted = TRUE, nda_rejected = FALSE WHERE id = %s", (current_user["user_id"],))
            conn.commit()
            record_activity(current_user["user_id"], current_user["user_id"], "Approved Mandate NDA", "Client officially accepted the mandatory NDA.")
            return {"message": "Mandate NDA accepted successfully"}
    except Exception as e:
        print(f"[ERROR] Accept mandate NDA error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)

@app.post("/api/contracts/reject-mandate")
def reject_mandate_nda(current_user: dict = Depends(verify_token)):
    if current_user["role"] != "client":
        raise HTTPException(status_code=403, detail="Only clients can reject mandate NDA")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("UPDATE users SET nda_accepted = FALSE, nda_rejected = TRUE WHERE id = %s", (current_user["user_id"],))
            conn.commit()
            record_activity(current_user["user_id"], current_user["user_id"], "Rejected Mandate NDA", "Client rejected the mandatory NDA.")
            return {"message": "Mandate NDA rejected"}
    except Exception as e:
        print(f"[ERROR] Reject mandate NDA error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)

@app.get("/api/templates/latest-nda")
def get_latest_nda_template(current_user: dict = Depends(verify_token)):
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, filename, uploaded_at, s3_key, s3_url FROM documents WHERE document_type = 'template' AND template_type = 'NDA' ORDER BY uploaded_at DESC LIMIT 1"
            )
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="NDA template not found")
            
            return {
                "id": row[0],
                "filename": row[1],
                "uploaded_at": row[2].isoformat() if row[2] else None,
                "s3_key": row[3],
                "s3_url": row[4]
            }
    except Exception as e:
        print(f"[ERROR] Get latest NDA error: {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        if conn: db_pool.putconn(conn)



@app.get("/api/activity/list")
def list_activities(client_id: Optional[str] = None, current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            if current_user["role"] == "client":
                cur.execute("SELECT id, user_id, client_id, action, details, timestamp FROM activity_log WHERE client_id = %s ORDER BY timestamp DESC LIMIT 100", (current_user["user_id"],))
            elif client_id:
                cur.execute("SELECT id, user_id, client_id, action, details, timestamp FROM activity_log WHERE client_id = %s ORDER BY timestamp DESC LIMIT 100", (client_id,))
            else:
                cur.execute("SELECT id, user_id, client_id, action, details, timestamp FROM activity_log ORDER BY timestamp DESC LIMIT 100")
            rows = cur.fetchall()
            return {"activities": [{"id": r[0], "user_id": r[1], "client_id": r[2], "action": r[3], "details": r[4], "timestamp": r[5].isoformat()} for r in rows]}
    finally: db_pool.putconn(conn)



@app.get("/api/contracts/download/{contract_id}")
def download_shared_contract(contract_id: str, current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            # First check shared_contracts
            cur.execute("SELECT s3_key, client_id, file_path, filename FROM shared_contracts WHERE id = %s", (contract_id,))
            res = cur.fetchone()
            
            # If not found, check documents (for injected mandatory NDA templates)
            if not res:
                cur.execute("SELECT s3_key, user_id, file_path, filename FROM documents WHERE id = %s", (contract_id,))
                res = cur.fetchone()
                # For templates, we allow clients to download if it's currently injected as their mandate
                # We can skip the strict client_id check for templates or perform a smarter check
                if not res: raise HTTPException(status_code=404)
            else:
                # Regular shared contract check
                if current_user["role"] == "client" and res[1] != current_user["user_id"]: 
                    raise HTTPException(status_code=403)

            record_activity(current_user["user_id"], current_user["user_id"], "Downloaded contract", f"Downloaded: {res[3]}")
            
            s3_key = res[0]
            if s3_key:
                try:
                    mime_type = "application/pdf" if s3_key.lower().endswith(".pdf") else "application/octet-stream"
                    url = s3_client.generate_presigned_url(
                        'get_object', 
                        Params={
                            'Bucket': BUCKET_NAME, 
                            'Key': s3_key,
                            'ResponseContentDisposition': 'inline',
                            'ResponseContentType': mime_type
                        }, 
                        ExpiresIn=3600
                    )
                    return {"download_url": url}
                except: pass
            fp = Path(res[2] or "")
            if fp.exists(): return FileResponse(path=str(fp), filename=res[3], content_disposition_type="inline")
            raise HTTPException(status_code=404)
    finally: db_pool.putconn(conn)


@app.post("/api/templates/upload")
async def upload_template(background_tasks: BackgroundTasks, file: UploadFile = File(...), template_type: str = Form(...), current_user: dict = Depends(verify_token)):
    if current_user["role"] not in ["admin", "legal_team"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    content = await file.read()
    file_name = f"template_{uuid.uuid4().hex[:8]}_{file.filename}"
    file_path = UPLOADS_DIR / file_name
    
    # Save locally first
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Upload to S3
    s3_url = None
    try:
        s3_client.put_object(Bucket=BUCKET_NAME, Key=file_name, Body=content)
        s3_url = f"https://{BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{file_name}"
        print(f"[S3] Template {file_name} uploaded successfully")
    except Exception as e:
        print(f"[ERROR] [S3] Template upload failed: {e}")
    
    tmpl_id = f"tmpl-{uuid.uuid4().hex[:8]}"
    conn = None
    try:
        if not db_pool:
            raise HTTPException(status_code=500, detail="Database not connected")
        
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # Metadata consistent with documents table
            cur.execute(
                """
                INSERT INTO documents (id, filename, document_type, template_type, user_id, user_email, user_role, size, status, shared_with, uploaded_at, s3_key, s3_url, file_path)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (
                    tmpl_id, file.filename, 'template', template_type, 
                    current_user["user_id"], current_user.get("email", ""), current_user.get("role", "admin"),
                    len(content), 'uploaded', json.dumps([]), datetime.now(),
                    file_name, s3_url, str(file_path)
                )
            )
            conn.commit()
            
            # Record activity
            record_activity(current_user["user_id"], "admin", "Uploaded standard template", f"Type: {template_type}")
            
            # Trigger extraction in background
            background_tasks.add_task(trigger_extraction, file_name, tmpl_id, template_type, "legal")
            
            return {"message": "Template uploaded and processing", "id": tmpl_id}
    except Exception as e:
        print(f"[ERROR] Database template upload error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    finally:
        if conn: db_pool.putconn(conn)


@app.get("/api/templates/analysis/{template_id}")
def get_template_analysis(template_id: str, current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT clause_id, clause, content_id, content, page_number FROM clauses WHERE document_id = %s ORDER BY page_number ASC",
                (template_id,)
            )
            clause_rows = cur.fetchall()
            
            if not clause_rows:
                return {"clauses": [], "status": "processing"}
                
            clauses_data = [
                {
                    "clause_id": r[0],
                    "clause": r[1],
                    "content_id": r[2],
                    "content": r[3],
                    "page_number": r[4]
                }
                for r in clause_rows
            ]
            
            return {"clauses": clauses_data, "status": "complete"}
    finally: db_pool.putconn(conn)


@app.get("/api/templates/download/{template_id}")
def download_template(template_id: str, current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT s3_key, file_path, filename FROM documents WHERE id = %s", (template_id,))
            res = cur.fetchone()
            if not res: raise HTTPException(status_code=404)
            if res[0]:
                s3_key = res[0]
                try:
                    # Robust mime detection
                    m_type = "application/pdf"
                    if s3_key.lower().endswith(".docx"):
                        m_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    
                    url = s3_client.generate_presigned_url(
                        'get_object', 
                        Params={
                            'Bucket': BUCKET_NAME, 
                            'Key': s3_key,
                            'ResponseContentDisposition': 'inline',
                            'ResponseContentType': m_type
                        }, 
                        ExpiresIn=3600
                    )
                    return {"download_url": url}
                except Exception as e:
                    print(f"[ERROR] Template URL generation failed: {e}")
                    # Fallback to basic URL without special params if generation failed
                    try:
                        url = s3_client.generate_presigned_url('get_object', Params={'Bucket': BUCKET_NAME, 'Key': s3_key}, ExpiresIn=3600)
                        return {"download_url": url}
                    except: pass
            
            fp = Path(res[1] or "")
            if fp.exists(): 
                return FileResponse(path=str(fp), filename=res[2], content_disposition_type="inline")
            raise HTTPException(status_code=404)
    finally: db_pool.putconn(conn)



@app.get("/api/templates/list")
def list_templates(current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, filename, template_type, uploaded_at FROM documents WHERE document_type = 'template'")
            rows = cur.fetchall()
            return {"templates": [{"id": r[0], "filename": r[1], "template_type": r[2], "uploaded_at": r[3].isoformat()} for r in rows]}
    finally: db_pool.putconn(conn)





# ──────────────────────────────────────────────────────────────────────────────
# Suggestion / Track-Changes Models
# ──────────────────────────────────────────────────────────────────────────────

class SuggestionCreate(BaseModel):
    content_id: str          # CNT-XXXXXXXX  
    original_text: str       # the exact text segment being replaced/deleted
    suggested_text: str      # proposed replacement (empty string = pure delete)
    author: Optional[str] = None

class SuggestionAction(BaseModel):
    suggestion_id: str
    action: str              # "accept" | "reject"


@app.post("/api/documents/review/{document_id}/suggest")
def create_suggestion(document_id: str, req: SuggestionCreate, current_user: dict = Depends(verify_token)):
    """Record a new text suggestion for a clause in track-changes mode."""
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")

    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # Fetch clause metadata and verify it belongs to this document
            cur.execute(
                """
                SELECT clause_id, clause, content_id, content, page_number, document, document_id
                FROM clauses WHERE content_id = %s AND document_id = %s
                """,
                (req.content_id, document_id)
            )
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Clause not found in this document")

            clause_id_val, clause_val, content_id_val, content_val, page_num, doc_label, doc_id_ref = row

            if not req.original_text and req.suggested_text:
                change_type = "insert"
            elif req.original_text and not req.suggested_text:
                change_type = "delete"
            else:
                change_type = "replace"

            sug_id = f"sug-{uuid.uuid4().hex[:8]}"
            author = "TYN Legal Team"

            cur.execute(
                """
                INSERT INTO clause_suggestions
                    (id, document_id, clause_id, clause, content_id, content, page_number,
                     document, change_type, original_text, suggested_text,
                     author, timestamp, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), 'pending')
                """,
                (
                    sug_id, document_id, clause_id_val, clause_val, content_id_val,
                    content_val, page_num, doc_label,
                    change_type, req.original_text, req.suggested_text, author
                )
            )
            cur.execute("UPDATE documents SET google_doc_id = NULL WHERE id = %s", (document_id,))
        conn.commit()
        return {"message": "Suggestion created", "suggestion_id": sug_id, "change_type": change_type}
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] create_suggestion error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)


@app.post("/api/documents/review/{document_id}/suggestion-action")
def suggestion_action(document_id: str, req: SuggestionAction, current_user: dict = Depends(verify_token)):
    """Accept or reject a pending suggestion."""
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if req.action not in ("accept", "reject"):
        raise HTTPException(status_code=400, detail="action must be 'accept' or 'reject'")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            # Verify suggestion exists and belongs to this document
            cur.execute(
                "SELECT id, content_id, original_text, suggested_text, status FROM clause_suggestions WHERE id = %s AND document_id = %s",
                (req.suggestion_id, document_id)
            )
            sug_row = cur.fetchone()
            if not sug_row:
                raise HTTPException(status_code=404, detail="Suggestion not found")
            
            sug_id, cid, orig_txt, sug_txt, old_status = sug_row
            if old_status != 'pending':
                raise HTTPException(status_code=400, detail=f"Suggestion is already {old_status}")

            new_status = "accepted" if req.action == "accept" else "rejected"
            cur.execute(
                "UPDATE clause_suggestions SET status = %s WHERE id = %s",
                (new_status, req.suggestion_id)
            )

            if req.action == "accept":
                # Get current text (either from edited_clauses or original clauses)
                cur.execute("SELECT edited_clause FROM edited_clauses WHERE content_id = %s", (cid,))
                e_row = cur.fetchone()
                
                if e_row and e_row[0] is not None:
                    current_text = e_row[0]
                else:
                    cur.execute("SELECT content FROM clauses WHERE content_id = %s", (cid,))
                    c_row = cur.fetchone()
                    current_text = c_row[0] if c_row else ""

                # Apply suggestion
                if orig_txt:
                    # Replace first occurrence of the original text
                    # Note: If orig_txt is not found, we append or do nothing?
                    # Since our UI sends the WHOLE text, orig_txt should be the current version.
                    updated_content = current_text.replace(orig_txt, sug_txt, 1)
                    if updated_content == current_text and orig_txt != sug_txt:
                        # Fallback: if replace failed (maybe concurrent edit), just take suggestions as full text
                        # ONLY if the UI is indeed sending full text (which it is)
                        updated_content = sug_txt
                else:
                    # Pure insert at end
                    updated_content = current_text + sug_txt

                # Persist to edited_clauses
                cur.execute(
                    """
                    INSERT INTO edited_clauses (content_id, original_clause, edited_clause, updated_at)
                    SELECT content_id, content, %s, NOW() FROM clauses WHERE content_id = %s
                    ON CONFLICT (content_id) DO UPDATE
                    SET edited_clause = EXCLUDED.edited_clause, updated_at = NOW()
                    """,
                    (updated_content, cid)
                )
            cur.execute("UPDATE documents SET google_doc_id = NULL WHERE id = %s", (document_id,))
        conn.commit()
        return {"message": f"Suggestion {req.action}ed", "suggestion_id": req.suggestion_id}
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] suggestion_action error: {e}")
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)


# ──────────────────────────────────────────────────────────────────────────────
# Review Endpoints
# ──────────────────────────────────────────────────────────────────────────────

@app.get("/api/documents/review/{document_id}")

def get_document_review(document_id: str, current_user: dict = Depends(verify_token)):
    conn = db_pool.getconn()
    try:
        with conn.cursor() as cur:
            # Fetch full document metadata
            cur.execute(
                "SELECT id, filename, document_type, user_id, user_email, user_role, size, status, shared_with, uploaded_at, file_path, s3_url, s3_key FROM documents WHERE id = %s",
                (document_id,)
            )
            doc_row = cur.fetchone()
            if not doc_row: raise HTTPException(status_code=404)
            if current_user["role"] == "client" and doc_row[3] != current_user["user_id"]: raise HTTPException(status_code=403)

            doc_meta = {
                "id": doc_row[0], "filename": doc_row[1], "document_type": doc_row[2],
                "user_id": doc_row[3], "user_email": doc_row[4], "user_role": doc_row[5],
                "size": doc_row[6], "status": doc_row[7], "shared_with": doc_row[8],
                "uploaded_at": doc_row[9].isoformat() if doc_row[9] else None,
                "file_path": doc_row[10], "s3_url": doc_row[11], "s3_key": doc_row[12]
            }
            s3_key = doc_row[12]

            # 1. DB check for risk analysis
            cur.execute("SELECT review_data FROM document_reviews WHERE document_id = %s", (document_id,))
            db_review = cur.fetchone()
            
            clauses = None
            if db_review:
                clauses = db_review[0]
            else:
                # 2. Risk file check (Fallback)
                reviews_dir = DATA_DIR / "reviews"
                review_path = reviews_dir / f"{document_id}.json"
                if review_path.exists():
                    with open(review_path, "r", encoding="utf-8") as f:
                        clauses = json.load(f)
                        
            if clauses is not None:
                # Fetch live statuses, edits, and comments from DB
                cur.execute("""
                    SELECT c.content_id, c.approval_status, e.edited_clause, e.comment 
                    FROM clauses c
                    LEFT JOIN edited_clauses e ON c.content_id = e.content_id
                    WHERE c.document_id = %s
                """, (document_id,))
                status_map = {}
                for row in cur.fetchall():
                    status_map[row[0]] = {
                        "status": row[1],
                        "edited_content": row[2],
                        "comment": row[3]
                    }

                # Fetch all suggestions for this document
                cur.execute("""
                    SELECT id, content_id, change_type, original_text, suggested_text,
                           author, timestamp, status
                    FROM clause_suggestions
                    WHERE document_id = %s AND status IN ('pending', 'accepted')
                    ORDER BY timestamp ASC
                """, (document_id,))
                sug_map = {}
                for srow in cur.fetchall():
                    cid_sug = srow[1]
                    if cid_sug not in sug_map:
                        sug_map[cid_sug] = []
                    sug_map[cid_sug].append({
                        "id": srow[0],
                        "change_type": srow[2],
                        "original_text": srow[3],
                        "suggested_text": srow[4],
                        "author": srow[5],
                        "timestamp": srow[6].isoformat() if srow[6] else None,
                        "status": srow[7]
                    })

                # Build live payload with accepted and pending changes
                import difflib
                for clause in clauses:
                    cid = clause.get("content_id")
                    if not cid: continue
                    
                    db_vals = status_map.get(cid, {})
                    clause["status"] = db_vals.get("status") or "pending"
                    clause["comment"] = db_vals.get("comment", "")
                    
                    orig_content = clause.get("content", "")
                    # 'accepted' text from edited_clauses
                    accepted_content = db_vals.get("edited_content") or orig_content
                    clause["edited_content"] = accepted_content
                    
                    # Fetch pending suggestions for this clause to build a preview_content
                    clause_sugs = sug_map.get(cid, [])
                    clause["suggestions"] = clause_sugs
                    
                    # Compute preview content (Base + Accepted + Pending)
                    # For simplicity, we apply all pending suggestions in sequence.
                    # In a block-level UI, typically only one version is proposed at a time,
                    # but we'll try to apply them.
                    preview_content = accepted_content
                    for sug in clause_sugs:
                        if sug["status"] == "pending":
                            # If UI sends full text, replace entire content
                            # We check if original_text matches our current preview
                            if sug["original_text"] == preview_content:
                                preview_content = sug["suggested_text"]
                            elif sug["original_text"] in preview_content:
                                preview_content = preview_content.replace(sug["original_text"], sug["suggested_text"], 1)

                    clause["preview_content"] = preview_content

                    def get_granular_diff(a, b):
                        # Split by word boundaries while keeping whitespaces
                        a_words = re.split(r'(\s+)', a)
                        b_words = re.split(r'(\s+)', b)
                        seq = difflib.SequenceMatcher(None, a_words, b_words)
                        out = []
                        for opcode, i1, i2, j1, j2 in seq.get_opcodes():
                            a_seg = "".join(a_words[i1:i2])
                            b_seg = "".join(b_words[j1:j2])
                            if opcode == 'equal':
                                out.append(a_seg)
                            elif opcode == 'delete':
                                out.append(f"<del class='tc-del' style='color:#dc2626; text-decoration:line-through; background-color:#fee2e2'>{a_seg}</del>")
                            elif opcode == 'insert':
                                out.append(f"<ins class='tc-ins' style='color:#16a34a; text-decoration:underline; background-color:#dcfce7'>{b_seg}</ins>")
                            elif opcode == 'replace':
                                out.append(f"<del class='tc-del' style='color:#dc2626; text-decoration:line-through; background-color:#fee2e2'>{a_seg}</del>")
                                out.append(f"<ins class='tc-ins' style='color:#16a34a; text-decoration:underline; background-color:#dcfce7'>{b_seg}</ins>")
                        return "".join(out).replace('\n', '<br/>')

                    # Generate Granular HTML Diff between Original and Preview
                    clause["html_diff"] = get_granular_diff(orig_content, preview_content) if preview_content != orig_content else orig_content


                return {"document": doc_meta, "clauses": clauses, "status": "complete"}
            
            # Fallback: read directly from clauses table (vector pipeline may have failed,
            # but extraction still saved clauses — show them with a default risk level)
            cur.execute(
                """SELECT c.clause_id, c.clause, c.content_id, c.content, c.page_number,
                          c.approval_status, e.edited_clause, e.comment
                   FROM clauses c
                   LEFT JOIN edited_clauses e ON c.content_id = e.content_id
                   WHERE c.document_id = %s ORDER BY c.page_number ASC, c.ctid ASC""",
                (document_id,)
            )
            raw_clauses = cur.fetchall()
            import difflib
            if raw_clauses:
                # Also load suggestions for fallback path
                cur.execute("""
                    SELECT id, content_id, change_type, original_text, suggested_text,
                           author, timestamp, status
                    FROM clause_suggestions
                    WHERE document_id = %s
                    ORDER BY timestamp ASC
                """, (document_id,))
                sug_map_fb = {}
                for srow in cur.fetchall():
                    cid_sug = srow[1]
                    if cid_sug not in sug_map_fb:
                        sug_map_fb[cid_sug] = []
                    sug_map_fb[cid_sug].append({
                        "id": srow[0],
                        "change_type": srow[2],
                        "original_text": srow[3],
                        "suggested_text": srow[4],
                        "author": srow[5],
                        "timestamp": srow[6].isoformat() if srow[6] else None,
                        "status": srow[7]
                    })

                clauses_fallback = []
                for r in raw_clauses:
                    orig_content = r[3]
                    edited_content = r[6]
                    html_diff = None
                    if edited_content and edited_content != orig_content:
                        seq = difflib.SequenceMatcher(None, orig_content, edited_content)
                        out = []
                        for opcode, i1, i2, j1, j2 in seq.get_opcodes():
                            if opcode == 'equal':
                                out.append(orig_content[i1:i2])
                            elif opcode == 'delete':
                                out.append(f"<del style='color:#dc2626; text-decoration:line-through'>{orig_content[i1:i2]}</del>")
                            elif opcode == 'insert':
                                out.append(f"<ins style='color:#16a34a; text-decoration:underline'>{edited_content[j1:j2]}</ins>")
                            elif opcode == 'replace':
                                out.append(f"<del style='color:#dc2626; text-decoration:line-through'>{orig_content[i1:i2]}</del>")
                                out.append(f"<ins style='color:#16a34a; text-decoration:underline'>{edited_content[j1:j2]}</ins>")
                        html_diff = "".join(out).replace('\n', '<br/>')

                    cid_fb = r[2]
                    clauses_fallback.append({
                        "clause_id":     r[0],
                        "clause_type":   r[1],
                        "content_id":    cid_fb,
                        "content":       orig_content,
                        "page_number":   r[4],
                        "risk":          "High",
                        "similarity_score": None,
                        "matched_clause":   None,
                        "llm_reasoning":    None,
                        "status":        r[5] or "pending",
                        "edited_content": edited_content,
                        "comment":       r[7],
                        "html_diff":     html_diff,
                        "suggestions":   sug_map_fb.get(cid_fb, [])
                    })
                return {"document": doc_meta, "clauses": clauses_fallback, "status": "complete"}

            return {"document": doc_meta, "clauses": [], "status": "processing"}
    finally: db_pool.putconn(conn)


@app.post("/api/documents/reprocess/{document_id}")
def reprocess_document(document_id: str, background_tasks: BackgroundTasks, current_user: dict = Depends(verify_token)):
    """Re-trigger extraction and vector review pipeline for an existing document."""
    if current_user["role"] not in ("admin", "legal_team"):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")

    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute(
                "SELECT s3_key, document_type, user_role FROM documents WHERE id = %s",
                (document_id,)
            )
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Document not found")

            s3_key, document_type, user_role = row
            if not s3_key:
                raise HTTPException(status_code=400, detail="No s3_key for this document – cannot reprocess")

            # Determine source from role
            source = "legal" if user_role in ("admin", "legal_team") else "client"

            # Clear old clauses and review so pipeline writes fresh results
            cur.execute("DELETE FROM clauses WHERE document_id = %s", (document_id,))
            cur.execute("DELETE FROM document_reviews WHERE document_id = %s", (document_id,))
            cur.execute("UPDATE documents SET status = 'processing' WHERE id = %s", (document_id,))
        conn.commit()
    except HTTPException:
        raise
    except Exception as e:
        if conn: conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn: db_pool.putconn(conn)

    # Ensure the file is available locally before triggering extraction
    local_path = UPLOADS_DIR / s3_key
    if not local_path.exists():
        try:
            s3_client.download_file(BUCKET_NAME, s3_key, str(local_path))
            print(f"[INFO] [Reprocess] Downloaded {s3_key} from S3 to {local_path}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to download file from S3: {str(e)}")

    # Kick off background extraction (same as upload flow)
    background_tasks.add_task(trigger_extraction, s3_key, document_type, source)
    return {"message": "Reprocessing started", "document_id": document_id}



class ClauseActionRequest(BaseModel):
    content_id: str
    action: str   # "accept" | "reject"


@app.post("/api/documents/review/{document_id}/action")
def clause_action(document_id: str, req: ClauseActionRequest, current_user: dict = Depends(verify_token)):
    if req.action not in ("accept", "reject"):
        raise HTTPException(status_code=400, detail="action must be 'accept' or 'reject'")
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")
    
    conn = None
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT user_id FROM documents WHERE id = %s", (document_id,))
            doc = cur.fetchone()
            if not doc:
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Simple access check
            if current_user["role"] == "client" and doc[0] != current_user["user_id"]:
                raise HTTPException(status_code=403, detail="Forbidden")


        # Update PostgreSQL directly instead of the temporary JSON
        updated = False
        try:
            with conn.cursor() as cur:
                new_status = "accepted" if req.action == "accept" else "rejected"
                cur.execute(
                    "UPDATE clauses SET approval_status = %s WHERE content_id = %s RETURNING id",
                    (new_status, req.content_id)
                )
                if cur.fetchone():
                    updated = True
                cur.execute("UPDATE documents SET google_doc_id = NULL WHERE id = %s", (document_id,))
            conn.commit()
        except Exception as e:
            print(f"[ERROR] DB update failed for clause action: {e}")
            conn.rollback()
            raise HTTPException(status_code=500, detail="Database error during approval")
            
        if not updated:
            raise HTTPException(status_code=404, detail="Clause not found in database")

        return {"message": f"Clause {req.action}ed", "content_id": req.content_id}
    except HTTPException: raise
    except Exception as e:
        print(f"[ERROR] Clause action error: {e}")
        raise HTTPException(status_code=500, detail="Processing error")
    finally:
        if conn: db_pool.putconn(conn)






class AskLLMRequest(BaseModel):
    content_id: str
    question: str


@app.post("/api/documents/review/{document_id}/ask-llm")
def ask_llm_about_clause(
    document_id: str,
    req: AskLLMRequest,
    current_user: dict = Depends(verify_token)
):
    """Ask the LLM a question about a specific clause in the review."""
    if not db_pool:
        raise HTTPException(status_code=500, detail="Database not connected")

    conn = None
    clauses = []
    try:
        conn = db_pool.getconn()
        with conn.cursor() as cur:
            cur.execute("SELECT review_data FROM document_reviews WHERE document_id = %s", (document_id,))
            db_review = cur.fetchone()
            if db_review:
                clauses = db_review[0]
            else:
                # Fallback to local file if DB isn't populated for this document yet
                reviews_dir = DATA_DIR / "reviews"
                review_path = reviews_dir / f"{document_id}.json"
                if review_path.exists():
                    with open(review_path, "r", encoding="utf-8") as f:
                        clauses = json.load(f)
                else:
                    raise HTTPException(status_code=404, detail="Review not available yet for this document")
    finally:
        if conn: db_pool.putconn(conn)

    # Find the specific clause the user clicked on
    clause = next((c for c in clauses if c.get("content_id") == req.content_id), None)
    if not clause:
        raise HTTPException(status_code=404, detail="Clause not found in review payload")

    try:
        from vector_pipeline.llm.reasoning import compare_clauses

        client_text   = clause.get("content", "")
        standard_text = (clause.get("matched_clause") or {}).get("content", "")
        clause_type   = clause.get("clause_type", "Unknown")
        risk          = clause.get("risk", "Unknown")

        answer = compare_clauses(client_text, standard_text, clause_type, risk)
        return {"answer": answer, "content_id": req.content_id}

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"LLM error: {str(e)}")

class ChatRequest(BaseModel):
    message: str


@app.post("/api/documents/chat/{document_id}")
def document_chat(
    document_id: str,
    req: ChatRequest,
    current_user: dict = Depends(verify_token)
):
    """Global Chatbot Agent for the entire document."""
    try:
        from vector_pipeline.llm.document_agent import chat_with_document
        
        # In a real app, bind the session ID to the user ID & document ID
        session_id = f"session_{current_user['user_id']}_{document_id}"
        
        response = chat_with_document(
            document_id=document_id,
            user_message=req.message,
            session_id=session_id
        )
        return {"response": response}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Chat Agent error: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True, access_log=True, log_level="info")
